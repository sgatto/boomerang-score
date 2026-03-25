
# rss8.py
# Vollständige Anwendung mit dynamischen Disziplinen, Inline-Edit, CSV/PDF-Export und Logo-Unterstützung.

import os
import sys

# Fix for [xcb] Unknown sequence number while appending request
if sys.platform.startswith("linux"):
    # Ensure XInitThreads is called before any X11-related library is loaded (like Tkinter).
    try:
        import ctypes
        x11 = ctypes.cdll.LoadLibrary("libX11.so.6")
        x11.XInitThreads()
    except Exception:
        pass
    # Setting this environment variable helps with X11/XCB sync issues
    os.environ["LIBXCB_ALLOW_SLOPPY_LOCK"] = "1"
    # Force X11 backend to avoid Wayland-related XCB issues
    os.environ["GDK_BACKEND"] = "x11"
    # Disable Xsynchronize to mitigate some race conditions in XCB
    os.environ["_X11_NO_XSYNCHRONIZE"] = "1"
    # QT_QPA_PLATFORM can also affect systems with mixed toolkits
    os.environ["QT_QPA_PLATFORM"] = "xcb"

from boomerang_score.core.scorer import compute_competition_ranks, ACC, AUS, MTA, END, FC, TC, TIMED

import csv
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

# Only import reportlab when needed
# from reportlab.lib import colors
# ...

# ==============================
# Ranking Helper Functions
# ==============================


# ==============================
# Discipline Configuration
# ==============================


# Helper sets
BASE_COLUMNS = ["name", "startnumber", "total", "overall_rank"]
EVENTS = ["ACC", "AUS", "MTA", "END", "FC", "TC", "TIMED"]
SORTED = ["StartNr", "Rank"]
# ==============================
# Main App
# ==============================
class ScoreTableApp(tk.Tk):
    """
    Dynamic competition list with discipline selection:
      - Base: Name | Start Number | Total Points | Overall Rank
      - Per active discipline: RES_[XX] | PTS_[XX] | RANK_[XX]

    - Inline-Edit: Name, Start Number, and all active RES values
    - Sorting: via click, toggle, numeric for numbers
    - CSV-Export: exports currently visible columns (displaycolumns)
    - PDF (Full list): A4 landscape, active disciplines
    - PDF (Individual reports): A4, compact table per participant, logo on the right
    """

    def __init__(self):
        super().__init__()
        self.title("Scoring Table – Dynamic Disciplines")
        self.geometry("1450x720")

        # Data storage: Dict[iid] -> Row-Dict (flat)
        # Fields: "name", "startnumber", "total", "overall_rank",
        #         for each discipline: "{code}_res", "{code}_pts", "{code}_rank"
        self.data = {}

        # Discipline status + entry fields
        self.disc_state = {d.code: tk.BooleanVar(value=d.default_active) for d in DISCIPLINES}
        self.disc_entries = {}  # code -> tk.Entry (Add area)

        # Tree/Columns
        self.tree = None
        self.all_columns = []        # complete column list (including invisible)
        self.display_columns = []    # currently visible columns
        self.column_visibility = {}  # key -> bool
        self.sort_state = {}         # Column -> ascending?

        # Inline-Editor
        self._edit_entry = None
        self._edit_iid_col = None

        # Logo + Title
        self.logo_path = None

        # Font configuration
        self.style = ttk.Style(self)
        self.font_main = ("Arial", 11)
        self.font_bold = ("Arial", 11, "bold")
        self.font_title = ("Arial", 16, "bold")
        
        # Apply global font via style
        self.style.configure(".", font=self.font_main)
        self.style.configure("Treeview", font=self.font_main)
        self.style.configure("Treeview.Heading", font=self.font_bold)
        
        # Ensure it also works for standard tk widgets if any
        self.option_add("*Font", self.font_main)

        self._build_ui()
        self._rebuild_dynamic_ui_and_tree()

    # =========================
    # UI Setup
    # =========================
    def _build_ui(self):
        # Menu
        menubar = tk.Menu(self)
        self.config(menu=menubar)
        view_menu = tk.Menu(menubar, tearoff=False)
        self.menu_view = view_menu
        menubar.add_cascade(label="View", menu=view_menu)
        view_menu.add_command(label="Manage columns …", command=self._open_columns_dialog)

        # Title + Logo
        frm_title = ttk.Frame(self, padding=(10, 6))
        frm_title.pack(fill="x")

        ttk.Label(frm_title, text="Competition Title:").grid(row=0, column=0, sticky="w")
        self.ent_title = ttk.Entry(frm_title, width=60)
        self.ent_title.grid(row=0, column=1, sticky="w", padx=(6, 12))
        self.ent_title.insert(0, "My Competition")

        self.lbl_title_display = ttk.Label(frm_title, text="My Competition", font=self.font_title)
        self.lbl_title_display.grid(row=1, column=0, columnspan=5, sticky="w", pady=(6, 2))

        def update_title(*_):
            self.lbl_title_display.config(text=self.ent_title.get())
        self.ent_title.bind("<KeyRelease>", update_title)

        ttk.Label(frm_title, text="Logo:").grid(row=0, column=2, sticky="e")
        self.lbl_logo_name = ttk.Label(frm_title, text="(no logo selected)")
        self.lbl_logo_name.grid(row=0, column=3, sticky="w", padx=(6, 6))
        ttk.Button(frm_title, text="Choose logo…", command=self.on_choose_logo).grid(row=0, column=4, sticky="w")

        # Discipline Checkboxes
        frm_disc = ttk.LabelFrame(self, text="Disciplines", padding=(10, 8))
        frm_disc.pack(fill="x", padx=10, pady=(0, 6))

        for idx, d in enumerate(DISCIPLINES):
            cb = ttk.Checkbutton(frm_disc, text=d.label, variable=self.disc_state[d.code],
                                 command=self._on_toggle_disciplines)
            cb.grid(row=0, column=idx, sticky="w", padx=(0, 12))

        # Input Area (Name, Start Number, dynamic discipline entries + buttons)
        frm_input = ttk.Frame(self, padding=(10, 8))
        frm_input.pack(fill="x")

        ttk.Label(frm_input, text="Name:").grid(row=0, column=0, sticky="w")
        self.ent_name = ttk.Entry(frm_input, width=20)
        self.ent_name.grid(row=0, column=1, sticky="w", padx=(6, 12))

        ttk.Label(frm_input, text="Start Number:").grid(row=0, column=2, sticky="w")
        self.ent_startnr = ttk.Entry(frm_input, width=10)
        self.ent_startnr.grid(row=0, column=3, sticky="w", padx=(6, 12))

        # Container for dynamic discipline fields
        self.frm_dyn_inputs = ttk.Frame(frm_input)
        self.frm_dyn_inputs.grid(row=0, column=4, sticky="w")

        # Move buttons to their own frame
        frm_buttons = ttk.Frame(frm_input)
        frm_buttons.grid(row=1, column=0, columnspan=105, sticky="w", pady=(8, 0))

        self.btn_add = ttk.Button(frm_buttons, text="Add line", command=self.on_add_row)
        self.btn_add.grid(row=0, column=0, padx=(0, 12))
        
        ttk.Button(frm_buttons, text="save CSV", command=self.export_csv).grid(row=0, column=1, padx=(0, 12))
        ttk.Button(frm_buttons, text="save PDF", command=self.export_pdf).grid(row=0, column=2, padx=(0, 12))
        ttk.Button(frm_buttons, text="Overall awards (PDF/DOCX)", command=self.export_individual_reports).grid(row=0, column=3, padx=(0, 12))

        # Scorsheet options in seperate frame
        frm_scoresheet = ttk.LabelFrame(self, text="Scoresheet", padding=(10, 8))
        frm_scoresheet.pack(fill="x", padx=10, pady=(0, 6))
        # frm_scoresheet.grid(row=2, column=0, columnspan=105, sticky="w", pady=(8, 0))

        ttk.Label(frm_scoresheet, text="Number of circles:").grid(row=0, column=2, sticky="w")
        self.ent_circle = ttk.Entry(frm_scoresheet, width=10)
        self.ent_circle.grid(row=0, column=3, sticky="w", padx=(6, 12))
        self.ent_circle.insert(0, "2")

        # Dropdown für Event
        ttk.Label(frm_scoresheet, text="Event:").grid(row=0, column=4, sticky="w", padx=(20, 6))

        self.event_var = tk.StringVar()
        self.event_var.set(EVENTS[0])  # Standardwert = erstes Element ("ACC")

        self.event_dropdown = ttk.OptionMenu(
            frm_scoresheet,
            self.event_var,
            self.event_var.get(),
            *EVENTS
        )
        self.event_dropdown.grid(row=0, column=5, sticky="w")
        # scoresheet_event = self.event_var.get()

        # Dropdown für Sortierung
        ttk.Label(frm_scoresheet, text="Sorted by").grid(row=0, column=6, sticky="w", padx=(20, 6))

        self.sheetsort_var = tk.StringVar()
        self.sheetsort_var.set(SORTED[0])  # Standardwert = erstes Element ("ACC")

        self.event_dropdown = ttk.OptionMenu(
            frm_scoresheet,
            self.sheetsort_var,
            self.sheetsort_var.get(),
            *SORTED
        )
        self.event_dropdown.grid(row=0, column=7, sticky="w")
        self.btn_add = ttk.Button(frm_scoresheet, text="print scoresheets", command=self.export_scoresheet)
        self.btn_add.grid(row=0, column=8, padx=(0, 12))


        for c in range(103):
            frm_input.grid_columnconfigure(c, weight=0)
        frm_input.grid_columnconfigure(103, weight=1)

        # Table (built dynamically)
        self.frm_table = ttk.Frame(self, padding=(10, 6))
        self.frm_table.pack(fill="both", expand=True)

    # =========================
    # Choose Logo
    # =========================
    def on_choose_logo(self):
        path = filedialog.askopenfilename(
            title="Select Logo File",
            filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.bmp;*.gif"), ("All files", "*.*")],
        )
        if not path:
            return
        self.logo_path = path
        self.lbl_logo_name.config(text=os.path.basename(path))

    # =========================
    # Toggle Disciplines
    # =========================
    def _on_toggle_disciplines(self):
        # Rebuild UI and table, recalculate ranks/total
        self._rebuild_dynamic_ui_and_tree()

    # =========================
    # Rebuild Dynamic UI & Tree
    # =========================
    def _rebuild_dynamic_ui_and_tree(self):
        # 1) Re-create dynamic entry fields
        for w in self.frm_dyn_inputs.winfo_children():
            w.destroy()
        self.disc_entries.clear()

        col = 0
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                ttk.Label(self.frm_dyn_inputs, text=f"{d.label} (Res):").grid(row=0, column=col, sticky="w", padx=(0, 4))
                ent = ttk.Entry(self.frm_dyn_inputs, width=8)
                ent.grid(row=0, column=col+1, sticky="w", padx=(0, 12))
                self.disc_entries[d.code] = ent
                col += 2

        # 2) Rebuild table (columns depend on active disciplines)
        #    We destroy/re-create Treeview, preserving order & data.
        old_children = []
        if self.tree is not None:
            old_children = list(self.tree.get_children(""))
            # Remember order
        for w in self.frm_table.winfo_children():
            w.destroy()
        self._build_tree()

        # 3) Re-insert existing data
        #    Order remains as stored in self.data (via Insert)
        for iid in self.data.keys():
            # Create item in tree, then update values
            new_iid = self.tree.insert("", "end", iid=iid, values=[""] * len(self.all_columns))
            self._update_tree_row(iid)

        # 4) Recalculate ranks/total as disciplines have changed
        self._recalc_ranks_and_update()

    # =========================
    # Create Tree Dynamically
    # =========================
    def _build_tree(self):
        # Assemble columns
        self.all_columns = []
        self.column_headers = {}
        self.column_widths = {}
        self.column_anchors = {}
        self.numeric_columns = set()

        # Base Columns
        base_defs = [
            ("name", "Name", 200, "w", False),
            ("startnumber", "Start No.", 80, "center", True),
            ("total", "Total Points", 120, "center", True),
            ("overall_rank", "Overall Rank", 100, "center", True),
        ]
        for key, hdr, w, anc, isnum in base_defs:
            self.all_columns.append(key)
            self.column_headers[key] = hdr
            self.column_widths[key] = w
            self.column_anchors[key] = anc
            if isnum:
                self.numeric_columns.add(key)

        # Discipline Columns (only active)
        for d in DISCIPLINES:
            if not self.disc_state[d.code].get():
                continue
            # Result
            key_e = f"{d.code}_res"
            self.all_columns.append(key_e)
            self.column_headers[key_e] = f"{d.label} Res"
            self.column_widths[key_e] = 90
            self.column_anchors[key_e] = "center"
            self.numeric_columns.add(key_e)
            # Points
            key_p = f"{d.code}_pts"
            self.all_columns.append(key_p)
            self.column_headers[key_p] = f"{d.label} Pts"
            self.column_widths[key_p] = 90
            self.column_anchors[key_p] = "center"
            self.numeric_columns.add(key_p)
            # Rank
            key_r = f"{d.code}_rank"
            self.all_columns.append(key_r)
            self.column_headers[key_r] = f"{d.label} Rank"
            self.column_widths[key_r] = 80
            self.column_anchors[key_r] = "center"
            self.numeric_columns.add(key_r)

        # Initialize/preserve visibility
        new_visibility = {}
        for key in self.all_columns:
            # Preserve known setting, otherwise default visible
            new_visibility[key] = self.column_visibility.get(key, True)
        self.column_visibility = new_visibility
        self.display_columns = [c for c in self.all_columns if self.column_visibility.get(c, True)]

        # Tree + Scrollbar
        self.tree = ttk.Treeview(
            self.frm_table,
            columns=self.all_columns,
            show="headings",
            selectmode="browse",
            height=20,
            displaycolumns=self.display_columns,
        )
        yscroll = ttk.Scrollbar(self.frm_table, orient="vertical", command=self.tree.yview)
        xscroll = ttk.Scrollbar(self.frm_table, orient="horizontal", command=self.tree.xview)
        self.tree.configure(xscroll=xscroll.set)
        self.tree.pack(side="bottom", fill="both", expand=True)
        xscroll.pack(side="bottom", fill="x")

        # Spalten konfigurieren
        for col in self.all_columns:
            self.tree.heading(col, text=self.column_headers[col], command=lambda c=col: self.on_sort_column(c))
            self.tree.column(col, width=self.column_widths[col], anchor=self.column_anchors[col], stretch=False)

        # Events: Inline-Edit
        self.tree.bind("<Double-1>", self.on_tree_double_click)
        self.bind("<Return>", self._commit_inline_edit)
        self.bind("<KP_Enter>", self._commit_inline_edit)
        self.bind("<Escape>", self._cancel_inline_edit)

    # =========================
    # Column Dialog (show/hide)
    # =========================
    def _open_columns_dialog(self):
        dlg = tk.Toplevel(self)
        dlg.title("Show/Hide Columns")
        dlg.transient(self)
        dlg.grab_set()
        frm = ttk.Frame(dlg, padding=10)
        frm.pack(fill="both", expand=True)

        vars_map = {}
        row = 0
        ttk.Label(frm, text="Visible Columns (only currently available):", font=self.font_bold).grid(row=row, column=0, sticky="w")
        row += 1
        for col in self.all_columns:
            v = tk.BooleanVar(value=self.column_visibility.get(col, True))
            vars_map[col] = v
            cb = ttk.Checkbutton(frm, text=self.column_headers[col], variable=v)
            cb.grid(row=row, column=0, sticky="w")
            row += 1

        btns = ttk.Frame(frm)
        btns.grid(row=row, column=0, pady=(8, 0), sticky="e")
        ttk.Button(btns, text="Cancel", command=dlg.destroy).pack(side="right", padx=(6, 0))
        def apply_and_close():
            # Apply visibility
            for col, v in vars_map.items():
                self.column_visibility[col] = bool(v.get())
            self.display_columns = [c for c in self.all_columns if self.column_visibility.get(c, True)]
            try:
                self.tree["displaycolumns"] = self.display_columns
            except Exception:
                pass
            dlg.destroy()
        ttk.Button(btns, text="Apply", command=apply_and_close).pack(side="right")

    # =========================
    # Parser/Formatter
    # =========================
    def _parse_float(self, s, allow_empty=True):
        s = (s or "").strip()
        if s == "":
            return None if allow_empty else 0.0
        return float(s.replace(",", "."))

    def _parse_int(self, s, allow_empty=True):
        s = (s or "").strip()
        if s == "":
            return None if allow_empty else 0
        return int(s)

    def _format_number(self, v):
        if v is None:
            return ""
        try:
            f = float(v)
        except (TypeError, ValueError):
            return str(v)
        if abs(f - int(f)) < 1e-9:
            return str(int(f))
        return f"{f:.2f}"

    def _next_free_startnr(self):
        used = {int(r.get("startnumber")) for r in self.data.values() if r.get("startnumber") is not None}
        n = 1
        while n in used:
            n += 1
        return n

    # =========================
    # Add Row
    # =========================
    def on_add_row(self):
        name = self.ent_name.get().strip()
        if not name:
            messagebox.showwarning("Missing Input", "Please enter a name.")
            return

        try:
            startnr = self._parse_int(self.ent_startnr.get())
        except ValueError:
            messagebox.showwarning("Invalid Start Number", "Start number must be an integer.")
            return
        if startnr is None or any(row.get("startnumber") == startnr for row in self.data.values()):
            startnr = self._next_free_startnr()

        # Capture values per active discipline
        disc_values = {}
        for d in DISCIPLINES:
            ent = self.disc_entries.get(d.code)
            if ent is None:
                disc_values[d.code] = None
            else:
                try:
                    v = self._parse_float(ent.get())
                except ValueError:
                    messagebox.showwarning("Invalid Input", f"{d.label}: Result must be a number.")
                    return
                disc_values[d.code] = v

        # Create new row
        iid = self.tree.insert("", "end", values=[""] * len(self.all_columns))
        row = {"name": name, "startnumber": startnr, "total": None, "overall_rank": None}
        # Initialize discipline fields
        for d in DISCIPLINES:
            row[f"{d.code}_res"] = float(disc_values[d.code]) if disc_values[d.code] is not None else 0.0
            row[f"{d.code}_pts"] = None
            row[f"{d.code}_rank"] = None

        self.data[iid] = row

        self._recalc_row(iid)
        self._update_tree_row(iid)
        self._recalc_ranks_and_update()

        # Clear inputs (name may remain)
        self.ent_startnr.delete(0, "end")
        for ent in self.disc_entries.values():
            ent.delete(0, "end")

    # =========================
    # Row Calculation/Display
    # =========================
    def _recalc_row(self, iid):
        row = self.data.get(iid)
        if not row:
            return
        # Points per discipline
        total = 0.0
        for d in DISCIPLINES:
            res = row.get(f"{d.code}_res") or 0.0
            pts = d.points_func(res)
            row[f"{d.code}_pts"] = float(pts)
            # Sum only over active disciplines
            if self.disc_state[d.code].get():
                total += float(pts)
        row["total"] = total

    def _update_tree_row(self, iid):
        row = self.data[iid]
        values = []
        for col in self.all_columns:
            if col in ("name",):
                values.append(row["name"])
            elif col in ("startnumber", "total", "overall_rank"):
                values.append(self._format_number(row.get(col)))
            elif col.endswith("_res") or col.endswith("_pts") or col.endswith("_rank"):
                values.append(self._format_number(row.get(col)))
            else:
                values.append("")
        self.tree.item(iid, values=values)

    def _recalc_ranks_and_update(self):
        # Discipline ranks (only active disciplines) by points
        for d in DISCIPLINES:
            if not self.disc_state[d.code].get():
                # Inactive discipline: clear rank
                for iid in self.data:
                    self.data[iid][f"{d.code}_rank"] = None
                continue
            if d.code == "fc":
                items = [(iid, self.data[iid].get(f"{d.code}_pts")) for iid in self.data]
            else:
                items = [(iid, self.data[iid].get(f"{d.code}_res")) for iid in self.data]
            ranks = compute_competition_ranks(items)
            for iid in self.data:
                self.data[iid][f"{d.code}_rank"] = ranks.get(iid)

        # Overall rank by total points
        items_total = [(iid, self.data[iid].get("total")) for iid in self.data]
        ranks_total = compute_competition_ranks(items_total)
        for iid in self.data:
            self.data[iid]["overall_rank"] = ranks_total.get(iid)
            self._update_tree_row(iid)

    # =========================
    # Inline-Editing
    # =========================
    def on_tree_double_click(self, event):
        region = self.tree.identify("region", event.x, event.y)
        if region != "cell":
            return
        rowid = self.tree.identify_row(event.y)
        colid = self.tree.identify_column(event.x)  # '#1', '#2', ...
        if not rowid or not colid:
            return

        col_index = int(colid.replace("#", "")) - 1
        col_key = self.tree["displaycolumns"][col_index]  # visible column
        # Editable: name, startnumber, *_res (of active ones!)
        editable = {"name", "startnumber"}
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                editable.add(f"{d.code}_res")

        if col_key not in editable:
            return

        bbox = self.tree.bbox(rowid, colid)
        if not bbox:
            return
        x, y, w, h = bbox

        current_text = self.tree.item(rowid, "values")[col_index]

        self._cancel_inline_edit()
        self._edit_entry = ttk.Entry(self.tree)
        self._edit_entry.insert(0, current_text)
        self._edit_entry.select_range(0, "end")
        self._edit_entry.focus_set()
        self._edit_entry.place(x=x, y=y, width=w, height=h)
        self._edit_iid_col = (rowid, col_key)
        self._edit_entry.bind("<FocusOut>", self._commit_inline_edit)

    def _commit_inline_edit(self, event=None):
        if not self._edit_entry or not self._edit_iid_col:
            return
        iid, col_key = self._edit_iid_col
        new_text = self._edit_entry.get()
        self._edit_entry.destroy()
        self._edit_entry = None
        self._edit_iid_col = None

        # Validate & apply
        if col_key == "name":
            if new_text.strip() == "":
                messagebox.showwarning("Invalid Name", "The name cannot be empty.")
                return
            self.data[iid]["name"] = new_text.strip()
            self._update_tree_row(iid)
            return

        if col_key == "startnumber":
            try:
                new_sn = int(new_text)
            except ValueError:
                messagebox.showwarning("Invalid Start Number", "The start number must be an integer.")
                return
            for oid in self.data:
                if oid != iid and self.data[oid].get("startnumber") == new_sn:
                    messagebox.showwarning("Duplicate Start Number", f"Start number {new_sn} is already assigned.")
                    return
            self.data[iid]["startnumber"] = new_sn
            self._update_tree_row(iid)
            return

        # Discipline result?
        if col_key.endswith("_res"):
            try:
                new_val = self._parse_float(new_text, allow_empty=False)
            except ValueError:
                messagebox.showwarning("Invalid Input", "Please enter a number.")
                return
            self.data[iid][col_key] = float(new_val)
            # Recalculate row + update ranks
            self._recalc_row(iid)
            self._update_tree_row(iid)
            self._recalc_ranks_and_update()
            return

    def _cancel_inline_edit(self, event=None):
        if self._edit_entry:
            self._edit_entry.destroy()
        self._edit_entry = None
        self._edit_iid_col = None

    # =========================
    # Sorting
    # =========================
    def _is_numeric_column(self, key):
        return key in self.numeric_columns

    def on_sort_column(self, col):
        asc = self.sort_state.get(col, True)
        asc = not asc
        self.sort_state[col] = asc
        self._apply_sort(col, asc)
        self._last_sort_col = col

    def _apply_sort(self, col, ascending=True):
        children = list(self.tree.get_children(""))

        def get_val(iid):
            row = self.data.get(iid, {})
            v = row.get(col)
            if v is None and col in ("name",):
                return str(row.get("name") or "")
            if v is None:
                try:
                    idx = self.tree["displaycolumns"].index(col)
                    val = self.tree.item(iid, "values")[idx]
                except Exception:
                    val = ""
                if self._is_numeric_column(col):
                    try:
                        return float(str(val).replace(",", "."))
                    except ValueError:
                        return float("-inf")
                return str(val)
            if self._is_numeric_column(col):
                try:
                    return float(v)
                except (TypeError, ValueError):
                    return float("-inf")
            return str(v)

        sorted_children = sorted(children, key=get_val, reverse=not ascending)
        for idx, iid in enumerate(sorted_children):
            self.tree.move(iid, "", idx)

    # =========================
    # Export CSV
    # =========================
    def export_csv(self):
        filename = filedialog.asksaveasfilename(
            title="Save CSV",
            defaultextension=".csv",
            filetypes=[("CSV file", "*.csv")]
        )
        if not filename:
            return

        # Export visible columns in current order
        cols = list(self.tree["displaycolumns"])
        headers = [self.column_headers[c] for c in cols]

        with open(filename, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f, delimiter=";")
            w.writerow(headers)
            for iid in self.tree.get_children():
                row = self.data[iid]
                out = []
                for c in cols:
                    out.append(row[c] if c in ("name",) else self._format_number(row.get(c)))
                w.writerow(out)

        messagebox.showinfo("Export Successful", f"The CSV has been saved:\n{filename}")

    # =========================
    # Export PDF Full List
    # =========================
    def export_pdf(self):
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        filename = filedialog.asksaveasfilename(
            title="Save PDF",
            defaultextension=".pdf",
            filetypes=[("PDF file", "*.pdf")]
        )
        if not filename:
            return

        from reportlab.lib.pagesizes import landscape
        doc = SimpleDocTemplate(
            filename,
            pagesize=landscape(A4),
            leftMargin=15*mm,
            rightMargin=15*mm,
            topMargin=15*mm,
            bottomMargin=15*mm,
        )

        styles = getSampleStyleSheet()
        story = []

        title_text = self.ent_title.get().strip() or "Competition"
        story.append(Paragraph(title_text, styles["Title"]))
        story.append(Spacer(1, 6))

        # Table: Base + active discipline columns (Res/Pts/Rank)
        headers = ["Name", "Start No.", "Total", "Overall Rank"]
        col_keys = ["name", "startnumber", "total", "overall_rank"]
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                headers += [f"{d.label} Res", f"{d.label} Pts", f"{d.label} Rank"]
                col_keys += [f"{d.code}_res", f"{d.code}_pts", f"{d.code}_rank"]

        data_rows = []
        for iid in self.tree.get_children():
            r = self.data[iid]
            row_vals = []
            for k in col_keys:
                if k == "name":
                    row_vals.append(str(r["name"]))
                else:
                    row_vals.append(self._format_number(r.get(k)))
            data_rows.append(row_vals)

        table_data = [headers] + data_rows

        # Column widths heuristic
        # Base: 60 + 18 + 22 + 24 = 124mm, rest for disciplines; per discipline ~ (22+22+20)=64mm
        col_widths = []
        base_widths = [45*mm, 11*mm, 11*mm, 11*mm]
        col_widths.extend(base_widths)
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                col_widths.extend([11*mm, 11*mm, 9*mm])

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 5),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),

            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 6),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("ALIGN", (0, 1), (0, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),

            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ]))
        story.append(tbl)

        try:
            doc.build(story)
        except Exception as e:
            messagebox.showerror("PDF Error", f"An error occurred while creating the PDF:\n{e}")
            return

        messagebox.showinfo("Export Successful", f"The PDF has been saved:\n{filename}")

    # =========================
    # Export: Individual Reports (A4, compact discipline table, logo)
    # =========================
    def export_individual_reports(self):
        # --- Bibliotheken prüfen ---
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        except Exception:
            pdf_available = False
        else:
            pdf_available = True

        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except Exception:
            docx_available = False
        else:
            docx_available = True

        if not pdf_available and not docx_available:
            messagebox.showerror(
                "Fehlende Pakete",
                "Weder ReportLab (PDF) noch python-docx (Word) sind installiert.\n"
                "Installiere mindestens eines davon:\n\n"
                "pip install reportlab\n"
                "pip install python-docx"
            )
            return

        if not self.data:
            messagebox.showwarning("No Data", "There are no participants.")
            return

        # Update ranks for current discipline selection
        self._recalc_ranks_and_update()

        # Sortierung
        entries = sorted(
            self.data.items(),
            key=lambda kv: ((kv[1].get("gesamtrang") or 10 ** 9), str(kv[1].get("name") or ""))
        )

        # --- Dateidialog: Nutzer entscheidet Format ---
        out_file = filedialog.asksaveasfilename(
            title="Bericht speichern",
            defaultextension=".pdf",
            filetypes=[
                ("PDF-Datei", "*.pdf"),
                ("Word-Dokument", "*.docx")
            ]
        )
        if not out_file:
            return

        # --- Format bestimmen ---
        is_pdf = out_file.lower().endswith(".pdf")
        is_docx = out_file.lower().endswith(".docx")

        title_text = self.ent_title.get().strip() or "Wettbewerb"

        # ---------------------------------------------------------
        # 1) WORD EXPORT
        # ---------------------------------------------------------
        if is_docx:
            if not docx_available:
                messagebox.showerror("Fehler", "python-docx ist nicht installiert.")
                return

            doc = Document()

            def add_logo(document):
                if not self.logo_path:
                    return
                try:
                    document.add_picture(self.logo_path, width=Inches(2.0))
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

            for idx, (iid, row) in enumerate(entries):
                h = doc.add_heading(title_text, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER

                h2 = doc.add_heading("Overall award", level=2)
                h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

                add_logo(doc)
                doc.add_paragraph("")

                # Teilnehmerkopf
                table = doc.add_table(rows=0, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                def add_info(label, value):
                    r = table.add_row().cells
                    r[0].text = label
                    r[1].text = str(value)

                add_info("Name:", row.get("name") or "")
                add_info("Gesamtpunkte:", self._format_number(row.get("gesamt")))
                add_info("Gesamtrang:", self._format_number(row.get("gesamtrang")))

                doc.add_paragraph("")

                # Disziplin-Tabelle
                disc_table = doc.add_table(rows=1, cols=4)
                disc_table.style = "Table Grid"
                hdr = disc_table.rows[0].cells
                hdr[0].text = "Disziplin"
                hdr[1].text = "Ergebnis"
                hdr[2].text = "Punkte"
                hdr[3].text = "Rang"

                for d in DISCIPLINES:
                    if not self.disc_state[d.code].get():
                        continue
                    row_cells = disc_table.add_row().cells
                    row_cells[0].text = d.label
                    row_cells[1].text = self._format_number(row.get(f"{d.code}_erg"))
                    row_cells[2].text = self._format_number(row.get(f"{d.code}_pkt"))
                    row_cells[3].text = self._format_number(row.get(f"{d.code}_rang"))

                if idx < len(entries) - 1:
                    doc.add_page_break()

            try:
                doc.save(out_file)
            except Exception as e:
                messagebox.showerror("Word-Fehler", f"Fehler beim Speichern:\n{e}")
                return

            messagebox.showinfo("Export erfolgreich", f"Word-Dokument gespeichert:\n{out_file}")
            return

        # ---------------------------------------------------------
        # 2) PDF EXPORT (dein bestehender Code)
        # ---------------------------------------------------------
        if is_pdf:
            if not pdf_available:
                messagebox.showerror("Fehler", "ReportLab ist nicht installiert.")
                return

            styles = getSampleStyleSheet()
            title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=20,
                                         leading=24, spaceAfter=6)
            h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14,
                                      spaceBefore=6, spaceAfter=6, alignment=1)
            label_style = ParagraphStyle("Label", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=11)
            text_style = ParagraphStyle("Text", parent=styles["Normal"], fontSize=11)

            def make_logo():
                if not self.logo_path:
                    return None
                try:
                    img = Image(self.logo_path)
                    max_w, max_h = 160 * mm, 160 * mm
                    iw, ih = img.imageWidth, img.imageHeight
                    scale = min(max_w / iw, max_h / ih)
                    img.drawWidth = iw * scale
                    img.drawHeight = ih * scale
                    return img
                except Exception:
                    return None

            title_text = self.ent_title.get().strip() or "Competition"

            def build_story_for_row(row):
                story = []
                logo = make_logo()
                story.append(Paragraph(title_text, title_style))
                story.append(Spacer(1, 6))
                story.append(Paragraph('<para alignment="center">Overall award</para>', h2_style))
                story.append(Spacer(1, 6))

                if logo:
                    logo.hAlign = "CENTER"
                    story.append(logo)

                story.append(Spacer(1, 20))

                # Teilnehmerkopf
                info_tbl = Table([
                    [Paragraph("Name:", label_style), Paragraph(str(row.get("name") or ""), text_style)],
                    # [Paragraph("Startnummer:", label_style), Paragraph(self._format_number(row.get("startnummer")), text_style)],
                    [Paragraph("Gesamtpunkte:", label_style),
                     Paragraph(self._format_number(row.get("gesamt")), text_style)],
                    [Paragraph("Gesamtrang:", label_style),
                     Paragraph(self._format_number(row.get("gesamtrang")), text_style)]
                ], colWidths=[40 * mm, None])
                #    info_tbl.setStyle(TableStyle([
                #        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                #        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                #        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                #    ]))

                info_tbl.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                ]))
                info_tbl.hAlign = "CENTER"

                story.append(info_tbl)
                story.append(Spacer(1, 30))

                # Kompakte Disziplin-Tabelle (nur aktive)
                tbl_headers = ["Disziplin", "Ergebnis", "Punkte", "Rang"]
                tbl_rows = []
                for d in DISCIPLINES:
                    if not self.disc_state[d.code].get():
                        continue
                    tbl_rows.append([
                        d.label,
                        self._format_number(row.get(f"{d.code}_erg")),
                        self._format_number(row.get(f"{d.code}_pkt")),
                        self._format_number(row.get(f"{d.code}_rang")),
                    ])
                table_data = [tbl_headers] + tbl_rows
                disc_tbl = Table(table_data, colWidths=[28 * mm, 28 * mm, 28 * mm, 22 * mm])
                disc_tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 6),

                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 1), (-1, -1), 10),
                    ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                    ("ALIGN", (0, 1), (0, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),

                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ]))
                story.append(disc_tbl)
                return story

            from reportlab.platypus import PageBreak
            doc = SimpleDocTemplate(out_file, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=16 * mm,
                                    bottomMargin=16 * mm)
            story = []
            for idx, (iid, row) in enumerate(entries):
                story.extend(build_story_for_row(row))
                if idx < len(entries) - 1:
                    story.append(PageBreak())
            try:
                doc.build(story)
            except Exception as e:
                messagebox.showerror("PDF-Fehler", f"Beim Erstellen des PDFs ist ein Fehler aufgetreten:\n{e}")
                return

            messagebox.showinfo("Export erfolgreich", f"PDF gespeichert:\n{out_file}")
            return

    # =========================
    # Export PDF Gesamtliste
    # =========================
    def export_scoresheet(self):
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
        except Exception:
            messagebox.showerror(
                "ReportLab fehlt",
                "Für den PDF-Export wird das Paket 'reportlab' benötigt.\n"
                "Installiere es mit:\n\n    pip install reportlab"
            )
            return

        filename = filedialog.asksaveasfilename(
            title="PDF speichern",
            defaultextension=".pdf",
            filetypes=[("PDF-Datei", "*.pdf")]
        )
        if not filename:
            return

        doc = SimpleDocTemplate(
            filename,
            pagesize=landscape(A4),
            leftMargin=15 * mm,
            rightMargin=15 * mm,
            topMargin=15 * mm,
            bottomMargin=15 * mm,
        )

        styles = getSampleStyleSheet()
        story = []

        title_text = self.ent_title.get().strip() or "Wettbewerb"
        story.append(Paragraph(title_text, styles["Title"]))
        story.append(Spacer(1, 6))

        # Tabelle: Basis + aktive Disziplinspalten (Erg/Pkt/Rang)
        headers = headers = ["startnr", "Name", "Wurf 1", "Wurf 2", "Wurf 3", "Wurf 4", "Wurf 5", "Gesamt"]
        col_keys = ["startnummer", "name", "gesamt", "gesamtrang"]
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                headers += [f"{d.label} Erg", f"{d.label} Pkt", f"{d.label} Rang"]
                col_keys += [f"{d.code}_erg", f"{d.code}_pkt", f"{d.code}_rang"]

        data_rows = []
        for iid in self.tree.get_children():
            r = self.data[iid]
            row_vals = []
            for k in col_keys:
                if k == "name":
                    row_vals.append(str(r["name"]))
                else:
                    row_vals.append(self._format_number(r.get(k)))
            data_rows.append(row_vals)

        table_data = [headers] + data_rows

        # Spaltenbreiten heuristisch
        # Basis: 60 + 18 + 22 + 24 = 124mm, Rest auf Disziplinen; pro Disziplin ~ (22+22+20)=64mm
        col_widths = []
        base_widths = [45 * mm, 11 * mm, 11 * mm, 11 * mm]
        col_widths.extend(base_widths)
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                col_widths.extend([11 * mm, 11 * mm, 9 * mm])

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 5),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),

            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 6),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("ALIGN", (0, 1), (0, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),

            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ]))
        story.append(tbl)

        try:
            doc.build(story)
        except Exception as e:
            messagebox.showerror("PDF Error", f"An error occurred while creating the PDF:\n{e}")
            return
        messagebox.showinfo("Export Successful", f"The PDF has been saved:\n{filename}")

    # =========================
    # Header/Start/Helper Functions
    # =========================
    def _is_active(self, code):
        return bool(self.disc_state[code].get())


# ==============================
# Starting Point
# ==============================
if __name__ == "__main__":
    app = ScoreTableApp()
    app.mainloop()

DISCIPLINES = [
    ACC,
    AUS,
    MTA,
    END,
    FC,
    TC,
    TIMED,
]
