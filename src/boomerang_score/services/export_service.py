"""Export functionality for competitions."""

import csv
from typing import Optional


def format_number(value) -> str:
    """Format a numeric value for display."""
    if value is None:
        return ""
    try:
        f = float(value)
    except (TypeError, ValueError):
        return str(value)
    if abs(f - int(f)) < 1e-9:
        return str(int(f))
    return f"{f:.2f}"


class ExportService:
    """Service for exporting competition data."""

    def __init__(self, competition, disciplines):
        """
        Initialize export service.

        Args:
            competition: Competition instance
            disciplines: List of discipline definitions
        """
        self.competition = competition
        self.disciplines = {d.code: d for d in disciplines}

    def export_csv(self, filename: str, visible_columns: list[str],
                   column_headers: dict[str, str], participant_order: list[str],
                   include_header: bool = False):
        """
        Export competition data to CSV.

        Args:
            filename: Output file path
            visible_columns: List of column keys to export
            column_headers: Dict mapping column keys to display names
            participant_order: List of participant IDs in display order
            include_header: Whether to include title and date header
        """
        import datetime
        headers = [column_headers[c] for c in visible_columns]

        with open(filename, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f, delimiter=";")

            if include_header:
                writer.writerow(["Titel", self.competition.title])
                writer.writerow(["Datum", datetime.date.today().isoformat()])
                writer.writerow([])

            writer.writerow(headers)

            for participant_id in participant_order:
                participant = self.competition.get_participant(participant_id)
                if not participant:
                    continue

                row = []
                for col in visible_columns:
                    if col == "name":
                        row.append(participant.name)
                    elif col == "startnumber":
                        row.append(format_number(participant.startnumber))
                    elif col == "total":
                        row.append(format_number(participant.total_points))
                    elif col == "overall_rank":
                        row.append(format_number(participant.overall_rank))
                    elif col.endswith("_res"):
                        disc_code = col[:-4]
                        row.append(format_number(participant.get_result(disc_code)))
                    elif col.endswith("_pts"):
                        disc_code = col[:-4]
                        row.append(format_number(participant.get_points(disc_code)))
                    elif col.endswith("_rank"):
                        disc_code = col[:-5]
                        row.append(format_number(participant.get_rank(disc_code)))
                    else:
                        row.append("")

                writer.writerow(row)

    def auto_save(self, all_columns: list[str], column_headers: dict[str, str], participant_order: list[str]):
        """
        Automatically save the competition to a CSV file.

        Args:
            all_columns: List of all column keys
            column_headers: Dict mapping column keys to display names
            participant_order: List of participant IDs in order
        """
        import os
        title = self.competition.title.strip() or "Wettbewerb"
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = os.path.join(os.getcwd(), safe_title + ".csv")

        try:
            self.export_csv(filename, all_columns, column_headers, participant_order, include_header=True)
        except Exception:
            # Silent failure for auto-save as in original code
            pass

    def load_csv(self, filename: str) -> dict:
        """
        Load competition data from a CSV file.

        Args:
            filename: Path to the CSV file

        Returns:
            Dict containing 'title' and 'participants' (list of dicts)
        """
        try:
            with open(filename, newline="", encoding="utf-8") as f:
                reader = csv.reader(f, delimiter=";")
                rows = list(reader)
        except Exception as e:
            raise ValueError(f"Could not read CSV file: {e}")

        if len(rows) < 4:
            raise ValueError("CSV file is too short or has invalid format.")

        title = "My Competition"
        if len(rows[0]) >= 2 and rows[0][0] == "Titel":
            title = rows[0][1]

        # Row 1: Date (ignored)
        # Row 2: Empty (ignored)
        # Row 3: Headers
        headers = rows[3]
        if not headers:
            raise ValueError("CSV file is missing headers.")

        participants_data = []
        for row in rows[4:]:
            if row:
                participants_data.append(dict(zip(headers, row)))

        return {
            "title": title,
            "participants": participants_data
        }

    def export_pdf_full_list(self, filename: str, participant_order: list[str]):
        """
        Export full competition list to PDF.

        Args:
            filename: Output file path
            participant_order: List of participant IDs in display order
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(
            filename,
            pagesize=landscape(A4),
            leftMargin=15*mm,
            rightMargin=15*mm,
            topMargin=15*mm,
            bottomMargin=15*mm,
        )

        styles = getSampleStyleSheet()
        story = []

        # Title
        title_text = self.competition.title or "Competition"
        story.append(Paragraph(title_text, styles["Title"]))
        story.append(Spacer(1, 6))

        # Build table headers and column keys
        headers = ["Name", "Start No.", "Total", "Overall Rank"]
        col_keys = ["name", "startnumber", "total", "overall_rank"]

        # Add active discipline columns
        active_disciplines = [
            (code, self.disciplines[code])
            for code in sorted(self.competition.active_disciplines)
            if code in self.disciplines
        ]

        for disc_code, disc in active_disciplines:
            headers += [f"{disc.label} Res", f"{disc.label} Pts", f"{disc.label} Rank"]
            col_keys += [f"{disc_code}_res", f"{disc_code}_pts", f"{disc_code}_rank"]

        # Build data rows
        data_rows = []
        for participant_id in participant_order:
            participant = self.competition.get_participant(participant_id)
            if not participant:
                continue

            row_vals = []
            for key in col_keys:
                if key == "name":
                    row_vals.append(participant.name)
                elif key == "startnumber":
                    row_vals.append(format_number(participant.startnumber))
                elif key == "total":
                    row_vals.append(format_number(participant.total_points))
                elif key == "overall_rank":
                    row_vals.append(format_number(participant.overall_rank))
                elif key.endswith("_res"):
                    disc_code = key[:-4]
                    row_vals.append(format_number(participant.get_result(disc_code)))
                elif key.endswith("_pts"):
                    disc_code = key[:-4]
                    row_vals.append(format_number(participant.get_points(disc_code)))
                elif key.endswith("_rank"):
                    disc_code = key[:-5]
                    row_vals.append(format_number(participant.get_rank(disc_code)))

            data_rows.append(row_vals)

        table_data = [headers] + data_rows

        # Column widths
        col_widths = [45*mm, 11*mm, 11*mm, 11*mm]  # Base columns
        for _ in active_disciplines:
            col_widths.extend([11*mm, 11*mm, 9*mm])  # Discipline columns

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 5),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 6),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("ALIGN", (0, 1), (0, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ]))
        story.append(tbl)

        doc.build(story)

    def export_individual_reports(self, filename: str, participant_order: list[str],
                                  logo_path: Optional[str] = None):
        """
        Export individual awards to PDF or DOCX.

        Args:
            filename: Output file path
            participant_order: List of participant IDs in display order
            logo_path: Optional path to logo image
        """
        is_pdf = filename.lower().endswith(".pdf")
        is_docx = filename.lower().endswith(".docx")

        if is_docx:
            self._export_individual_docx(filename, participant_order, logo_path)
        elif is_pdf:
            self._export_individual_pdf(filename, participant_order, logo_path)
        else:
            raise ValueError("Filename must end with .pdf or .docx")

    def export_scoresheet(self, filename: str, event: str, num_circles: int,
                          sort_method: str, participant_order: list[str]):
        """
        Export competition scoresheets to PDF or DOCX.

        Args:
            filename: Output file path
            event: Discipline code or event name
            num_circles: Number of circles to distribute participants into
            sort_method: How to sort and distribute ("StartNr" or "Rank")
            participant_order: Base participant order (by StartNr or Rank)
        """
        is_pdf = filename.lower().endswith(".pdf")
        is_docx = filename.lower().endswith(".docx")

        # Distribute entries to circles
        circles_data = self._distribute_entries_to_circles(participant_order, num_circles, sort_method)

        # Get scoresheet columns configuration
        header_rows, num_empty_cols, col_widths = self._get_scoresheet_columns(event)

        if is_docx:
            self._export_scoresheet_docx(filename, event, circles_data, header_rows, col_widths)
        elif is_pdf:
            self._export_scoresheet_pdf(filename, event, circles_data, header_rows, num_empty_cols, col_widths)
        else:
            raise ValueError("Filename must end with .pdf or .docx")

    def _get_scoresheet_columns(self, event: str):
        """Get columns configuration for a specific event scoresheet."""
        from reportlab.lib.units import mm
        event = event.upper()

        if event == "ACC":
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 11)] + ["Result"]
            num_empty_cols = 10
            col_widths = [12*mm, 40*mm] + [12*mm] * 10 + [15*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "AUS":
            main_headers = ["Startnr", "Name"]
            sub_headers = ["", ""]
            for i in range(1, 6):
                main_headers.extend([f"Throw {i}", "", "", ""])
                sub_headers.extend(["Dist", "Catch", "Acc", "To"])
            main_headers.append("Result")
            sub_headers.append("")
            num_empty_cols = 20
            col_widths = [12*mm, 35*mm] + [10*mm] * 20 + [15*mm]
            return [main_headers, sub_headers], num_empty_cols, col_widths
        elif event == "END":
            distances = [5, 10, 15, 20, 25, 30, 35, 40, 45, 50, 55, 60, 65, 70, 75, 80]
            headers = ["Startnr", "Name"] + [str(d) for d in distances] + ["Result"]
            num_empty_cols = 16
            col_widths = [12*mm, 35*mm] + [10*mm] * 16 + [15*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "FC":
            headers = ["Startnr", "Name", "Round 1", "Round 2", "Result"]
            num_empty_cols = 2
            col_widths = [12*mm, 40*mm] + [40*mm] * 2 + [40*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "MTA":
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 6)] + ["Result"]
            num_empty_cols = 5
            col_widths = [12*mm, 40*mm] + [20*mm] * 5 + [20*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "TC":
            main_headers = ["Startnr", "Name"] + ["Throw 1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "", "Doubl.1", "", "2", "", "3", "", "4", "", "5", ""] + ["Result"]
            sub_headers = ["", ""] + ["Left-\nhand\nclean", "Right-\nhand\nclean", "2 hand\nbehind\nback", "2 hand\nunder\nthe-leg", "Eagle\ncatch", "Hacky\ncatch", "Tunnel\ncatch", "1 hand\nbehind\nback", "1 hand\nunder\nthe leg", "Foot\ncatch", "Total\nSingle", "2 hand\nbehind-\nthe-back", "2 hand\nunder-\nthe-leg", "Left-\nhand\nclean", "Hacky\ncatch", "Right-\nhand\nclean", "Tunnel\ncatch", "1 hand\nbehind\nthe back", "1 hand\nunder\nthe leg", "Eagle\ncatch", "Foot\ncatch"] + [""]
            points_headers = ["", ""] + ["(3)", "(3)", "(4)", "(3)", "(4)", "(7)", "(5)", "(7)", "(6)", "(8)", "", "(4)", "(3)", "(3)", "(7)", "(3)", "(5)", "(7)", "(6)", "(4)", "(8)"] + [""]
            num_empty_cols = 21
            col_widths = [12*mm, 35*mm] + [9*mm] * 21 + [15*mm]
            return [main_headers, sub_headers, points_headers], num_empty_cols, col_widths
        elif event == "TAPIR":
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 6)] + ["Result"]
            num_empty_cols = 5
            col_widths = [12*mm, 40*mm] + [15*mm] * 5 + [15*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "TIMED":
            headers = ["Startnr", "Name", "Round 1", "Round 2", "Result"]
            num_empty_cols = 2
            col_widths = [12*mm, 40*mm] + [40*mm] * 2 + [40*mm]
            return [headers], num_empty_cols, col_widths
        else:
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 11)] + ["Result"]
            num_empty_cols = 10
            col_widths = [12*mm, 30*mm] + [12*mm] * 10 + [15*mm]
            return [headers], num_empty_cols, col_widths

    def _distribute_entries_to_circles(self, participant_order: list[str], num_circles: int, sort_method: str):
        """Distribute participants into circles."""
        circles_data = [[] for _ in range(num_circles)]

        if sort_method == "Rank":
            # Cyclic distribution for rank-based sorting
            for idx, p_id in enumerate(participant_order):
                circle_idx = idx % num_circles
                circles_data[circle_idx].append(p_id)
        else:
            # Sequential distribution for startnr-based sorting
            total_entries = len(participant_order)
            entries_per_circle = max(1, (total_entries + num_circles - 1) // num_circles)
            for idx, p_id in enumerate(participant_order):
                circle_idx = idx // entries_per_circle
                if circle_idx >= num_circles:
                    circle_idx = num_circles - 1
                circles_data[circle_idx].append(p_id)

        return circles_data

    def _export_scoresheet_pdf(self, filename, event, circles_data, header_rows, num_empty_cols, col_widths):
        """Export scoresheet to PDF using ReportLab."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(
            filename,
            pagesize=landscape(A4),
            leftMargin=10*mm,
            rightMargin=10*mm,
            topMargin=10*mm,
            bottomMargin=10*mm,
        )

        styles = getSampleStyleSheet()
        story = []
        title_text = self.competition.title or "Competition"

        for circle_idx, circle_p_ids in enumerate(circles_data, 1):
            if not circle_p_ids:
                continue
            if circle_idx > 1:
                story.append(PageBreak())

            story.append(Paragraph(f"{title_text} - {event} - Circle {circle_idx}", styles["Title"]))
            story.append(Spacer(1, 6))

            table_data = [list(r) for r in header_rows]

            for p_id in circle_p_ids:
                participant = self.competition.get_participant(p_id)
                if not participant:
                    continue
                row_vals = [str(participant.startnumber), participant.name]
                row_vals.extend([""] * (num_empty_cols + 1))
                table_data.append(row_vals)

            tbl = Table(table_data, colWidths=col_widths, repeatRows=len(header_rows))
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, len(header_rows)-1), colors.HexColor("#d3d3d3")),
                ("TEXTCOLOR", (0, 0), (-1, len(header_rows)-1), colors.black),
                ("ALIGN", (0, 0), (-1, len(header_rows)-1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, len(header_rows)-1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, len(header_rows)-1), 7),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (0, len(header_rows)), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, len(header_rows)), (-1, -1), 7),
                ("ALIGN", (0, len(header_rows)), (0, -1), "CENTER"),
                ("ALIGN", (1, len(header_rows)), (1, -1), "LEFT"),
                ("ALIGN", (2, len(header_rows)), (-1, -1), "CENTER"),
                ("ROWBACKGROUNDS", (0, len(header_rows)), (-1, -1), [colors.whitesmoke, colors.white]),
            ]))
            story.append(tbl)

        doc.build(story)

    def _export_scoresheet_docx(self, filename, event, circles_data, header_rows, col_widths):
        """Export scoresheet to DOCX."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        def set_cell_background(cell, fill):
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), fill)
            cell._element.get_or_add_tcPr().append(shading_elm)

        title_text = self.competition.title or "Competition"
        num_cols = len(header_rows[-1])

        for circle_idx, circle_p_ids in enumerate(circles_data, 1):
            if not circle_p_ids:
                continue
            if circle_idx > 1:
                doc.add_page_break()

            heading = doc.add_heading(f"{title_text} - {event} - Circle {circle_idx}", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            table = doc.add_table(rows=len(header_rows), cols=num_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for r_idx, header_row in enumerate(header_rows):
                hdr_cells = table.rows[r_idx].cells
                for c_idx, text in enumerate(header_row):
                    hdr_cells[c_idx].text = text
                    set_cell_background(hdr_cells[c_idx], "D3D3D3")
                    for p in hdr_cells[c_idx].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.bold = True
                            run.font.size = Pt(9)

            for p_id in circle_p_ids:
                participant = self.competition.get_participant(p_id)
                if not participant:
                    continue
                row_cells = table.add_row().cells
                row_cells[0].text = str(participant.startnumber)
                row_cells[1].text = participant.name
                for i in range(2, num_cols):
                    row_cells[i].text = ""

                for i, cell in enumerate(row_cells):
                    for p in cell.paragraphs:
                        if i == 0 or i == num_cols - 1:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif i == 1:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(filename)

    def _export_individual_pdf(self, filename: str, participant_order: list[str],
                               logo_path: Optional[str]):
        """Export individual reports to PDF."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                       Paragraph, Spacer, Image, PageBreak)
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ReportTitle", parent=styles["Title"],
            fontName="Helvetica-Bold", fontSize=20, leading=24, spaceAfter=6
        )
        h2_style = ParagraphStyle(
            "H2", parent=styles["Heading2"],
            fontName="Helvetica-Bold", fontSize=14,
            spaceBefore=6, spaceAfter=6, alignment=1
        )

        def make_logo():
            if not logo_path:
                return None
            try:
                img = Image(logo_path)
                max_w, max_h = 160 * mm, 160 * mm
                iw, ih = img.imageWidth, img.imageHeight
                scale = min(max_w / iw, max_h / ih)
                img.drawWidth = iw * scale
                img.drawHeight = ih * scale
                return img
            except Exception:
                return None

        doc = SimpleDocTemplate(
            filename, pagesize=A4,
            leftMargin=18*mm, rightMargin=18*mm,
            topMargin=16*mm, bottomMargin=16*mm
        )

        story = []
        title_text = self.competition.title or "Competition"

        for idx, participant_id in enumerate(participant_order):
            participant = self.competition.get_participant(participant_id)
            if not participant:
                continue

            # Title and header
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 6))
            story.append(Paragraph('<para alignment="center">Overall award</para>', h2_style))
            story.append(Spacer(1, 6))

            # Logo
            logo = make_logo()
            if logo:
                logo.hAlign = "CENTER"
                story.append(logo)
            story.append(Spacer(1, 20))

            # Participant info table
            label_style = ParagraphStyle("Label", parent=styles["Normal"],
                                        fontName="Helvetica-Bold", fontSize=11)
            text_style = ParagraphStyle("Text", parent=styles["Normal"], fontSize=11)

            info_tbl = Table([
                [Paragraph("Name:", label_style),
                 Paragraph(participant.name, text_style)],
                [Paragraph("Total Points:", label_style),
                 Paragraph(format_number(participant.total_points), text_style)],
                [Paragraph("Overall Rank:", label_style),
                 Paragraph(format_number(participant.overall_rank), text_style)]
            ], colWidths=[40*mm, None])
            info_tbl.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
            ]))
            info_tbl.hAlign = "CENTER"
            story.append(info_tbl)
            story.append(Spacer(1, 30))

            # Discipline results table
            tbl_headers = ["Discipline", "Result", "Points", "Rank"]
            tbl_rows = []

            for disc_code in sorted(self.competition.active_disciplines):
                if disc_code not in self.disciplines:
                    continue
                disc = self.disciplines[disc_code]
                tbl_rows.append([
                    disc.label,
                    format_number(participant.get_result(disc_code)),
                    format_number(participant.get_points(disc_code)),
                    format_number(participant.get_rank(disc_code)),
                ])

            table_data = [tbl_headers] + tbl_rows
            disc_tbl = Table(table_data, colWidths=[28*mm, 28*mm, 28*mm, 22*mm])
            disc_tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 1), (-1, -1), 10),
                ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                ("ALIGN", (0, 1), (0, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
            ]))
            story.append(disc_tbl)

            # Page break between participants
            if idx < len(participant_order) - 1:
                story.append(PageBreak())

        doc.build(story)

    def _export_individual_docx(self, filename: str, participant_order: list[str],
                                logo_path: Optional[str]):
        """Export individual reports to DOCX."""
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        title_text = self.competition.title or "Competition"

        for idx, participant_id in enumerate(participant_order):
            participant = self.competition.get_participant(participant_id)
            if not participant:
                continue

            # Title
            h = doc.add_heading(title_text, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            h2 = doc.add_heading("Overall award", level=2)
            h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Logo
            if logo_path:
                try:
                    doc.add_picture(logo_path, width=Inches(2.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

            doc.add_paragraph("")

            # Participant info
            table = doc.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            def add_info(label, value):
                r = table.add_row().cells
                r[0].text = label
                r[1].text = str(value)

            add_info("Name:", participant.name)
            add_info("Total Points:", format_number(participant.total_points))
            add_info("Overall Rank:", format_number(participant.overall_rank))

            doc.add_paragraph("")

            # Discipline table
            disc_table = doc.add_table(rows=1, cols=4)
            disc_table.style = "Table Grid"
            hdr = disc_table.rows[0].cells
            hdr[0].text = "Discipline"
            hdr[1].text = "Result"
            hdr[2].text = "Points"
            hdr[3].text = "Rank"

            for disc_code in sorted(self.competition.active_disciplines):
                if disc_code not in self.disciplines:
                    continue
                disc = self.disciplines[disc_code]
                row_cells = disc_table.add_row().cells
                row_cells[0].text = disc.label
                row_cells[1].text = format_number(participant.get_result(disc_code))
                row_cells[2].text = format_number(participant.get_points(disc_code))
                row_cells[3].text = format_number(participant.get_rank(disc_code))

            # Page break between participants
            if idx < len(participant_order) - 1:
                doc.add_page_break()

        doc.save(filename)
