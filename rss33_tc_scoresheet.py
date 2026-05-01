

# Vollständige Anwendung mit dynamischen Disziplinen, Inline-Edit, CSV/PDF-Export und Logo-Unterstützung.

import os
import math
import tkinter as tk
from tkinter import ttk, messagebox, filedialog

# ==============================
# Ranking-Hilfsfunktion
# ==============================
def compute_competition_ranks(items):
    """
    Standard-Wettbewerbsranking (1, 1, 3, 4, 4, 6, ...)
    items: Liste[(iid, value)], höherer Wert = besser
    Rückgabe: dict[iid] -> rank
    """
    norm = []
    for iid, val in items:
        try:
            v = float(val)
        except (TypeError, ValueError):
            v = float("-inf")
        norm.append((iid, v))
    norm.sort(key=lambda x: x[1], reverse=True)

    ranks = {}
    prev_val = None
    rank = 0
    count = 0
    for iid, v in norm:
        count += 1
        if prev_val is None or v != prev_val:
            rank = count
            prev_val = v
        ranks[iid] = rank
    return ranks


# ==============================
# Disziplin-Konfiguration
# ==============================
class Discipline:
    def __init__(self, code, label, default_active, points_func):
        self.code = code      # 'acc'
        self.label = label    # 'ACC'
        self.default_active = default_active
        self.points_func = points_func

def safe_div(numer, denom):
    try:
        denom = float(denom)
        if denom == 0:
            return 0.0
        return float(numer) / denom
    except (TypeError, ValueError):
        return 0.0

def _points_100(_erg):
    max_score_100 = 100
    if _erg < 0:
       _loc_points = -200
    elif _erg < 100:
       _loc_points = 500 * math.log10( 1 + 99 * (_erg / max_score_100))
    else:
       _loc_points = 1000
    return _loc_points
    
def _points_80(_erg):
    max_score_80 = 80
    if _erg < 0:
       _loc_points = -200
    else:
       _loc_points = 500 * math.log10( 1 + 99 * (_erg / max_score_80))
    return _loc_points

def _points_50(_erg):
    max_score_50 = 50
    if _erg < 0:
       _loc_points = -200
    elif _erg < max_score_50:
       _loc_points = 500 * math.log10( 1 + 99 * (_erg / max_score_50))
    else:
       _loc_points = 1000
    return _loc_points

def _points_fc(_erg):
    _max_time = 60.0
    if _erg == 0:
       _loc_points = 0
    elif _erg == 1:
       _loc_points = 387.26
    elif _erg == 2:
       _loc_points = 518.71
    elif _erg == 3:
       _loc_points = 600.01
    elif _erg == 4:
       _loc_points = 659.03
    elif _erg >= 75:
       _loc_points = 659.03
    elif _erg >= 5:
       _loc_points = 500 * math.log10( 1 + 99 * ( 15.00/_erg))
    else:
       _loc_points = -200
    return _loc_points
    
def _points_timed(_erg):
    _max_time = 60.0
    if _erg == 0:
       _loc_points = 0
    elif _erg == 1:
       _loc_points = 387.26
    elif _erg == 2:
       _loc_points = 518.71
    elif _erg == 3:
       _loc_points = 600.01
    elif _erg == 4:
       _loc_points = 659.03
    elif _erg >= 75:
       _loc_points = 659.03
    elif _erg > 5:
       _loc_points = 500 * math.log10( 1 + 99 * safe_div( 15.00/_erg))
    else:
       _loc_points = -200
    return _loc_points


def _points_tapir(_erg):
    try:
        return float(_erg) * 3
    except (TypeError, ValueError):
        return 0.0

DISCIPLINES = [
    Discipline("acc", "ACC",   True,  lambda e: _points_100(float(e)) ),
    Discipline("aus", "AUS",   True,  lambda e: _points_100(float(e)) ),
    Discipline("mta", "MTA",   True,  lambda e: _points_50(float(e)) ),
    Discipline("end", "END",   True,  lambda e: _points_80(float(e)) ),
    Discipline("fc",  "FC",    True,  lambda e: _points_fc(float(e)) ),
    Discipline("tc",  "TC",    True,  lambda e: _points_100(float(e)) ),
    Discipline("timed","TIMED",False, lambda e: _points_timed(float(e)) ),
    Discipline("tapir","TAPIR",True,  lambda e: _points_tapir(float(e)) ),
]

# Hilfssets
BASE_COLUMNS = ["name", "startnummer", "gesamt", "gesamtrang"]
EVENTS = ["ACC", "AUS", "MTA", "END", "FC", "TC", "TIMED", "TAPIR"]
SORTED = ["StartNr", "Rank"]
# ==============================
# Haupt-App
# ==============================
class ScoreTableApp(tk.Tk):
    """
    Dynamische Wettbewerbsliste mit Disziplin-Auswahl:
      - Basis: Name | Startnummer | Gesamtpunkte | Gesamtrang
      - je aktive Disziplin: ERG_[XX] | PKT_[XX] | RANG_[XX]

    - Inline-Edit: Name, Startnummer, und alle aktiven ERG-Werte
    - Sortierung: per Klick, Toggle, numerisch für Zahlen
    - CSV-Export: exportiert die aktuell sichtbaren Spalten (displaycolumns)
    - PDF (Gesamtliste): A4 quer, aktive Disziplinen
    - PDF (Einzelberichte): A4, kompakte Tabelle je Teilnehmer, Logo rechts
    """

    def __init__(self):
        super().__init__()
        self.title("Wertungstabelle – Dynamische Disziplinen")
        self.geometry("1450x720")

        # Datenhaltung: Dict[iid] -> Row-Dict (flat)
        # Felder: "name", "startnummer", "gesamt", "gesamtrang",
        #         für jede Disziplin: "{code}_erg", "{code}_pkt", "{code}_rang"
        self.data = {}

        # Disziplinstatus + Eingabefelder
        self.disc_state = {d.code: tk.BooleanVar(value=d.default_active) for d in DISCIPLINES}
        self.disc_entries = {}  # code -> tk.Entry (Add-Bereich)

        # Tree/Spalten
        self.tree = None
        self.all_columns = []        # komplette Spaltenliste (inkl. unsichtbare)
        self.display_columns = []    # aktuell sichtbare Spalten
        self.column_visibility = {}  # key -> bool
        self.sort_state = {}         # Spalte -> aufsteigend?

        # Inline-Editor
        self._edit_entry = None
        self._edit_iid_col = None

        # Logo + Titel
        self.logo_path = None

        self._build_ui()
        self._rebuild_dynamic_ui_and_tree()

    # =========================
    # UI-Aufbau mit Scrollbalken für ganzes Fenster
    # =========================
    def _build_ui(self):
        # Haupt-Container mit Scrollbalken für das gesamte Fenster
        self.main_canvas = tk.Canvas(self, highlightthickness=0)
        self.main_canvas.pack(side="left", fill="both", expand=True)

        # Scrollbalken für das Hauptfenster
        self.v_scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.main_canvas.yview)
        self.v_scrollbar.pack(side="right", fill="y")
        self.h_scrollbar = ttk.Scrollbar(self, orient="horizontal", command=self.main_canvas.xview)
        self.h_scrollbar.pack(side="bottom", fill="x")

        self.main_canvas.configure(yscrollcommand=self.v_scrollbar.set, xscrollcommand=self.h_scrollbar.set)

        # Frame im Canvas für alle Widgets
        self.main_frame = ttk.Frame(self.main_canvas)
        self.main_canvas.create_window((0, 0), window=self.main_frame, anchor="nw")

        # Event-Binding für Canvas-Scrolling
        self.main_frame.bind("<Configure>", self._on_frame_configure)
        self.main_canvas.bind("<MouseWheel>", self._on_mousewheel)
        self.main_canvas.bind("<Shift-MouseWheel>", self._on_shift_mousewheel)

        # Alle Widgets werden jetzt in self.main_frame gepackt statt in self
        # Menü
        menubar = tk.Menu(self)
        self.config(menu=menubar)
        view_menu = tk.Menu(menubar, tearoff=False)
        self.menu_view = view_menu
        menubar.add_cascade(label="Ansicht", menu=view_menu)
        view_menu.add_command(label="Spalten verwalten …", command=self._open_columns_dialog)

        # Alle Widgets werden jetzt in self.main_frame gepackt statt in self

        # Titel + Logo
        frm_title = ttk.Frame(self.main_frame, padding=(10, 6))
        frm_title.pack(fill="x")

        ttk.Label(frm_title, text="Wettbewerbstitel:").grid(row=0, column=0, sticky="w")
        self.ent_title = ttk.Entry(frm_title, width=60)
        self.ent_title.grid(row=0, column=1, sticky="w", padx=(6, 12))
        self.ent_title.insert(0, "Mein Wettbewerb")

        self.lbl_title_display = ttk.Label(frm_title, text="Mein Wettbewerb", font=("Arial", 16, "bold"))
        self.lbl_title_display.grid(row=1, column=0, columnspan=5, sticky="w", pady=(6, 2))

        def update_title(*_):
            self.lbl_title_display.config(text=self.ent_title.get())
        self.ent_title.bind("<KeyRelease>", update_title)

        ttk.Label(frm_title, text="Logo:").grid(row=0, column=2, sticky="e")
        self.lbl_logo_name = ttk.Label(frm_title, text="(kein Logo gewählt)")
        self.lbl_logo_name.grid(row=0, column=3, sticky="w", padx=(6, 6))
        ttk.Button(frm_title, text="Logo wählen…", command=self.on_choose_logo).grid(row=0, column=4, sticky="w")

        # Disziplin-Checkboxen
        frm_disc = ttk.LabelFrame(self.main_frame, text="Disziplinen", padding=(10, 8))
        frm_disc.pack(fill="x", padx=10, pady=(0, 6))

        for idx, d in enumerate(DISCIPLINES):
            cb = ttk.Checkbutton(frm_disc, text=d.label, variable=self.disc_state[d.code],
                                 command=self._on_toggle_disciplines)
            cb.grid(row=0, column=idx, sticky="w", padx=(0, 12))

        # Eingabe-Bereich (Name, Startnr, dynamische Disziplin-Eingaben + Buttons)
        frm_input = ttk.Frame(self.main_frame, padding=(10, 8))
        frm_input.pack(fill="x")

        ttk.Label(frm_input, text="Name:").grid(row=0, column=0, sticky="w")
        self.ent_name = ttk.Entry(frm_input, width=20)
        self.ent_name.grid(row=0, column=1, sticky="w", padx=(6, 12))

        ttk.Label(frm_input, text="Startnummer:").grid(row=0, column=2, sticky="w")
        self.ent_startnr = ttk.Entry(frm_input, width=10)
        self.ent_startnr.grid(row=0, column=3, sticky="w", padx=(6, 12))

        # Container für dynamische Disziplin-Eingaben
        self.frm_dyn_inputs = ttk.Frame(frm_input)
        self.frm_dyn_inputs.grid(row=0, column=4, sticky="w")

        # Buttons in eigenes Frame verschieben
        frm_buttons = ttk.Frame(frm_input)
        frm_buttons.grid(row=1, column=0, columnspan=105, sticky="w", pady=(8, 0))
        
        self.btn_add = ttk.Button(frm_buttons, text="Add line", command=self.on_add_row)
        self.btn_add.grid(row=0, column=0, padx=(0, 12))
        ttk.Button(frm_buttons, text="Delete line", command=self.on_delete_row).grid(row=0, column=1, padx=(0, 12))
        
        ttk.Button(frm_buttons, text="load CSV", command=self.load_csv).grid(row=0, column=2, padx=(0, 12))
        ttk.Button(frm_buttons, text="save CSV", command=self.export_csv).grid(row=0, column=3, padx=(0, 12))
        ttk.Button(frm_buttons, text="save PDF", command=self.export_pdf).grid(row=0, column=4, padx=(0, 12))
        ttk.Button(frm_buttons, text="Overall awards (PDF/DOCX)", command=self.export_individual_reports).grid(row=0, column=5, padx=(0, 12))
        
        # Scorsheet options in seperate frame
        frm_scoresheet = ttk.LabelFrame(self.main_frame, text="Scoresheet", padding=(10, 8))
        frm_scoresheet.pack(fill="x", padx=10, pady=(0, 6))
        #frm_scoresheet.grid(row=2, column=0, columnspan=105, sticky="w", pady=(8, 0))
        
        ttk.Label(frm_scoresheet, text="Number of circles:").grid(row=0, column=2, sticky="w")
        self.ent_circle = ttk.Entry(frm_scoresheet, width=10)
        self.ent_circle.grid(row=0, column=3, sticky="w", padx=(6, 12))
        self.ent_circle.insert(0, "2")
        
        # Dropdown für Event
        ttk.Label(frm_scoresheet, text="Event:").grid(row=0, column=4, sticky="w", padx=(20, 6))

        self.event_var = tk.StringVar()
        self.event_var.set(EVENTS[0])  # Standardwert = erstes Element ("ACC")

        self.event_dropdown = ttk.OptionMenu(
            frm_scoresheet,
            self.event_var,
            self.event_var.get(),
            *EVENTS
        )
        self.event_dropdown.grid(row=0, column=5, sticky="w")
        #scoresheet_event = self.event_var.get()
        
        # Dropdown für Sortierung
        ttk.Label(frm_scoresheet, text="Sorted by").grid(row=0, column=6, sticky="w", padx=(20, 6))

        self.sheetsort_var = tk.StringVar()
        self.sheetsort_var.set(SORTED[0])  # Standardwert = erstes Element ("ACC")

        self.event_dropdown = ttk.OptionMenu(
            frm_scoresheet,
            self.sheetsort_var,
            self.sheetsort_var.get(),
            *SORTED
        )
        self.event_dropdown.grid(row=0, column=7, sticky="w")
        self.btn_add = ttk.Button(frm_scoresheet, text="print scoresheets", command=self.export_scoresheet)
        self.btn_add.grid(row=0, column=8, padx=(0, 12))
        

        for c in range(103):
            frm_input.grid_columnconfigure(c, weight=0)
        frm_input.grid_columnconfigure(103, weight=1)

        # Tabelle (wird dynamisch aufgebaut)
        self.frm_table = ttk.Frame(self.main_frame, padding=(10, 6))
        self.frm_table.pack(fill="both", expand=True)

    # =========================
    # Scroll-Methoden für das Hauptfenster
    # =========================
    def _on_frame_configure(self, event=None):
        """Passt die Scroll-Region des Canvas an, wenn sich der Frame ändert"""
        self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all"))

    def _on_mousewheel(self, event):
        """Vertikales Scrolling mit Mausrad"""
        self.main_canvas.yview_scroll(int(-1*(event.delta/120)), "units")

    def _on_shift_mousewheel(self, event):
        """Horizontales Scrolling mit Shift+Mausrad"""
        self.main_canvas.xview_scroll(int(-1*(event.delta/120)), "units")

    # =========================
    # Logo wählen
    # =========================
    def on_choose_logo(self):
        path = filedialog.askopenfilename(
            title="Logo-Datei auswählen",
            filetypes=[("Bilddateien", "*.png;*.jpg;*.jpeg;*.bmp;*.gif"), ("Alle Dateien", "*.*")],
        )
        if not path:
            return
        self.logo_path = path
        self.lbl_logo_name.config(text=os.path.basename(path))

    # =========================
    # Disziplinen toggeln
    # =========================
    def _on_toggle_disciplines(self):
        # UI und Tabelle neu aufbauen, Ränge/Total neu berechnen
        self._rebuild_dynamic_ui_and_tree()

    # =========================
    # Dynamische UI & Tree aufbauen
    # =========================
    def _rebuild_dynamic_ui_and_tree(self):
        # 1) Dynamische Eingabefelder neu erstellen
        for w in self.frm_dyn_inputs.winfo_children():
            w.destroy()
        self.disc_entries.clear()

        col = 0
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                ttk.Label(self.frm_dyn_inputs, text=f"{d.label} (Erg):").grid(row=0, column=col, sticky="w", padx=(0, 4))
                ent = ttk.Entry(self.frm_dyn_inputs, width=8)
                ent.grid(row=0, column=col+1, sticky="w", padx=(0, 12))
                self.disc_entries[d.code] = ent
                col += 2

        # 2) Tabelle neu aufbauen (Spalten abhängig von aktiven Disziplinen)
        #    Wir zerstören/re-erstellen Treeview, übernehmen Reihenfolge & Daten.
        old_children = []
        if self.tree is not None:
            old_children = list(self.tree.get_children(""))
            # Reihenfolge merken
        for w in self.frm_table.winfo_children():
            w.destroy()
        self._build_tree()

        # 3) vorhandene Daten wieder einfügen
        #    Reihenfolge bleibt so, wie self.data gespeichert ist (durch Insert)
        for iid in self.data.keys():
            # Item in Tree anlegen, dann Werte aktualisieren
            new_iid = self.tree.insert("", "end", iid=iid, values=[""] * len(self.all_columns))
            self._update_tree_row(iid)

        # 4) Ränge/Total neu berechnen, da Disziplinen sich geändert haben
        self._recalc_ranks_and_update()

    # =========================
    # Tree dynamisch erzeugen
    # =========================
    def _build_tree(self):
        # Spalten zusammenstellen
        self.all_columns = []
        self.column_headers = {}
        self.column_widths = {}
        self.column_anchors = {}
        self.numeric_columns = set()

        # Basis-Spalten
        base_defs = [
            ("name", "Name", 200, "w", False),
            ("startnummer", "Startnr.", 80, "center", True),
            ("gesamt", "Gesamtpunkte", 120, "center", True),
            ("gesamtrang", "Gesamtrang", 100, "center", True),
        ]
        for key, hdr, w, anc, isnum in base_defs:
            self.all_columns.append(key)
            self.column_headers[key] = hdr
            self.column_widths[key] = w
            self.column_anchors[key] = anc
            if isnum:
                self.numeric_columns.add(key)

        # Disziplin-Spalten (nur aktive)
        for d in DISCIPLINES:
            if not self.disc_state[d.code].get():
                continue
            # Ergebnis
            key_e = f"{d.code}_erg"
            self.all_columns.append(key_e)
            self.column_headers[key_e] = f"{d.label} Erg"
            self.column_widths[key_e] = 90
            self.column_anchors[key_e] = "center"
            self.numeric_columns.add(key_e)
            # Punkte
            key_p = f"{d.code}_pkt"
            self.all_columns.append(key_p)
            self.column_headers[key_p] = f"{d.label} Pkt"
            self.column_widths[key_p] = 90
            self.column_anchors[key_p] = "center"
            self.numeric_columns.add(key_p)
            # Rang
            key_r = f"{d.code}_rang"
            self.all_columns.append(key_r)
            self.column_headers[key_r] = f"{d.label} Rang"
            self.column_widths[key_r] = 80
            self.column_anchors[key_r] = "center"
            self.numeric_columns.add(key_r)

        # Sichtbarkeit initialisieren/erhalten
        new_visibility = {}
        for key in self.all_columns:
            # bereits bekannte Einstellung übernehmen, sonst default sichtbar
            new_visibility[key] = self.column_visibility.get(key, True)
        self.column_visibility = new_visibility
        self.display_columns = [c for c in self.all_columns if self.column_visibility.get(c, True)]

        # Tree + Scrollbar - Scrollbalken innerhalb des Treeview
        self.style = ttk.Style(self)
        self.style.configure("Treeview.Heading", font=("Arial", 10, "bold"), background="#d9d9d9")
        self.style.configure("Treeview", rowheight=24, font=("Arial", 10), fieldbackground="#ffffff")
        self.style.map("Treeview", background=[('selected', '#ececec')], foreground=[('selected', '#000000')])

        self.tree = ttk.Treeview(
            self.frm_table,
            columns=self.all_columns,
            show="headings",
            selectmode="browse",
            height=20,
            displaycolumns=self.display_columns,
        )

        # Keine separaten Scrollbalken für Treeview - das ganze Fenster scrollt
        self.tree.pack(fill="both", expand=True)

        # Alternate row colors
        self.tree.tag_configure("oddrow", background="#F8F9FB")
        self.tree.tag_configure("evenrow", background="#FFFFFF")

        # Spalten konfigurieren
        for col in self.all_columns:
            self.tree.heading(col, text=self.column_headers[col], command=lambda c=col: self.on_sort_column(c))
            self.tree.column(col, width=self.column_widths[col], anchor=self.column_anchors[col], stretch=False)

        # Events: Inline-Edit
        self.tree.bind("<Double-1>", self.on_tree_double_click)
        self.bind("<Return>", self._commit_inline_edit)
        self.bind("<KP_Enter>", self._commit_inline_edit)
        self.bind("<Escape>", self._cancel_inline_edit)

    # =========================
    # Spalten-Dialog (ein/ausblenden)
    # =========================
    def _open_columns_dialog(self):
        dlg = tk.Toplevel(self)
        dlg.title("Spalten ein-/ausblenden")
        dlg.transient(self)
        dlg.grab_set()
        frm = ttk.Frame(dlg, padding=10)
        frm.pack(fill="both", expand=True)

        vars_map = {}
        row = 0
        ttk.Label(frm, text="Sichtbare Spalten (nur aktuell verfügbare):", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky="w")
        row += 1
        for col in self.all_columns:
            v = tk.BooleanVar(value=self.column_visibility.get(col, True))
            vars_map[col] = v
            cb = ttk.Checkbutton(frm, text=self.column_headers[col], variable=v)
            cb.grid(row=row, column=0, sticky="w")
            row += 1

        btns = ttk.Frame(frm)
        btns.grid(row=row, column=0, pady=(8, 0), sticky="e")
        ttk.Button(btns, text="Abbrechen", command=dlg.destroy).pack(side="right", padx=(6, 0))
        def apply_and_close():
            # Sichtbarkeit anwenden
            for col, v in vars_map.items():
                self.column_visibility[col] = bool(v.get())
            self.display_columns = [c for c in self.all_columns if self.column_visibility.get(c, True)]
            try:
                self.tree["displaycolumns"] = self.display_columns
            except Exception:
                pass
            dlg.destroy()
        ttk.Button(btns, text="Übernehmen", command=apply_and_close).pack(side="right")

    # =========================
    # Parser/Formatter
    # =========================
    def _parse_float(self, s, allow_empty=True):
        s = (s or "").strip()
        if s == "":
            return None if allow_empty else 0.0
        return float(s.replace(",", "."))

    def _parse_int(self, s, allow_empty=True):
        s = (s or "").strip()
        if s == "":
            return None if allow_empty else 0
        return int(s)

    def _format_number(self, v):
        if v is None:
            return ""
        try:
            f = float(v)
        except (TypeError, ValueError):
            return str(v)
        if abs(f - int(f)) < 1e-9:
            return str(int(f))
        return f"{f:.2f}"

    def _next_free_startnr(self):
        used = {int(r.get("startnummer")) for r in self.data.values() if r.get("startnummer") is not None}
        n = 1
        while n in used:
            n += 1
        return n

    # =========================
    # Zeile hinzufügen
    # =========================
    def on_add_row(self):
        name = self.ent_name.get().strip()
        if not name:
            messagebox.showwarning("Eingabe fehlt", "Bitte einen Namen eingeben.")
            return

        try:
            startnr = self._parse_int(self.ent_startnr.get())
        except ValueError:
            messagebox.showwarning("Ungültige Startnummer", "Startnummer muss eine ganze Zahl sein.")
            return
        if startnr is None or any(row.get("startnummer") == startnr for row in self.data.values()):
            startnr = self._next_free_startnr()

        # Werte je aktive Disziplin erfassen
        disc_values = {}
        for d in DISCIPLINES:
            ent = self.disc_entries.get(d.code)
            if ent is None:
                disc_values[d.code] = None
            else:
                try:
                    v = self._parse_float(ent.get())
                except ValueError:
                    messagebox.showwarning("Ungültige Eingabe", f"{d.label}: Ergebnis muss Zahl sein.")
                    return
                disc_values[d.code] = v

        # Neue Zeile anlegen
        iid = self.tree.insert("", "end", values=[""] * len(self.all_columns))
        row = {"name": name, "startnummer": startnr, "gesamt": None, "gesamtrang": None}
        # Disziplin-Felder initialisieren
        for d in DISCIPLINES:
            row[f"{d.code}_erg"] = float(disc_values[d.code]) if disc_values[d.code] is not None else 0.0
            row[f"{d.code}_pkt"] = None
            row[f"{d.code}_rang"] = None

        self.data[iid] = row

        self._recalc_row(iid)
        self._update_tree_row(iid)
        self._recalc_ranks_and_update()

        self._auto_save()

        # Eingaben leeren (Name darf stehen bleiben)
        self.ent_startnr.delete(0, "end")
        for ent in self.disc_entries.values():
            ent.delete(0, "end")

    # =========================
    # Zeile löschen
    # =========================
    def on_delete_row(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Auswahl fehlt", "Bitte wähle eine Zeile zum Löschen aus.")
            return

        if not messagebox.askyesno("Zeile löschen", "Soll die ausgewählte Zeile wirklich gelöscht werden?"):
            return

        for iid in selected:
            self.tree.delete(iid)
            self.data.pop(iid, None)

        self._recalc_ranks_and_update()
        self._auto_save()

    # =========================
    # Zeile rechnen/anzeigen
    # =========================
    def _recalc_row(self, iid):
        row = self.data.get(iid)
        if not row:
            return
        # Punkte je Disziplin
        total = 0.0
        for d in DISCIPLINES:
            erg = row.get(f"{d.code}_erg") or 0.0
            pkt = d.points_func(erg)
            row[f"{d.code}_pkt"] = float(pkt)
            # Summe nur über aktive Disziplinen
            if self.disc_state[d.code].get():
                total += float(pkt)
        row["gesamt"] = total

    def _update_tree_row(self, iid):
        row = self.data[iid]
        values = []
        for col in self.all_columns:
            if col in ("name",):
                values.append(row["name"])
            elif col in ("startnummer", "gesamt", "gesamtrang"):
                values.append(self._format_number(row.get(col)))
            elif col.endswith("_erg") or col.endswith("_pkt") or col.endswith("_rang"):
                values.append(self._format_number(row.get(col)))
            else:
                values.append("")

        # alternating row tags for better visual clarity
        children = list(self.tree.get_children(""))
        if iid in children:
            idx = children.index(iid)
        else:
            idx = len(children)

        tag = "evenrow" if idx % 2 == 0 else "oddrow"
        self.tree.item(iid, values=values, tags=(tag,))

    def _recalc_ranks_and_update(self):
        # Disziplin-Ränge (nur aktive Disziplinen) nach Punkten
        for d in DISCIPLINES:
            if not self.disc_state[d.code].get():
                # Inaktive Disziplin: Rang leeren
                for iid in self.data:
                    self.data[iid][f"{d.code}_rang"] = None
                continue
            if d.code == "fc":
                items = [(iid, self.data[iid].get(f"{d.code}_pkt")) for iid in self.data]
            else:
                items = [(iid, self.data[iid].get(f"{d.code}_erg")) for iid in self.data]
            ranks = compute_competition_ranks(items)
            for iid in self.data:
                self.data[iid][f"{d.code}_rang"] = ranks.get(iid)

        # Gesamtrang nach Gesamtpunkten
        items_total = [(iid, self.data[iid].get("gesamt")) for iid in self.data]
        ranks_total = compute_competition_ranks(items_total)
        for iid in self.data:
            self.data[iid]["gesamtrang"] = ranks_total.get(iid)
            self._update_tree_row(iid)

    # =========================
    # Inline-Editing
    # =========================
    def on_tree_double_click(self, event):
        region = self.tree.identify("region", event.x, event.y)
        if region != "cell":
            return
        rowid = self.tree.identify_row(event.y)
        colid = self.tree.identify_column(event.x)  # '#1', '#2', ...
        if not rowid or not colid:
            return

        col_index = int(colid.replace("#", "")) - 1
        col_key = self.tree["displaycolumns"][col_index]  # sichtbare Spalte
        # Editierbar: name, startnummer, *_erg (der aktiven!)
        editable = {"name", "startnummer"}
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                editable.add(f"{d.code}_erg")

        if col_key not in editable:
            return

        bbox = self.tree.bbox(rowid, colid)
        if not bbox:
            return
        x, y, w, h = bbox

        current_text = self.tree.item(rowid, "values")[col_index]

        self._cancel_inline_edit()
        self._edit_entry = ttk.Entry(self.tree)
        self._edit_entry.insert(0, current_text)
        self._edit_entry.select_range(0, "end")
        self._edit_entry.focus_set()
        self._edit_entry.place(x=x, y=y, width=w, height=h)
        self._edit_iid_col = (rowid, col_key)
        self._edit_entry.bind("<FocusOut>", self._commit_inline_edit)

    def _commit_inline_edit(self, event=None):
        if not self._edit_entry or not self._edit_iid_col:
            return
        iid, col_key = self._edit_iid_col
        new_text = self._edit_entry.get()
        self._edit_entry.destroy()
        self._edit_entry = None
        self._edit_iid_col = None

        # Validieren & übernehmen
        if col_key == "name":
            if new_text.strip() == "":
                messagebox.showwarning("Ungültiger Name", "Der Name darf nicht leer sein.")
                return
            self.data[iid]["name"] = new_text.strip()
            self._update_tree_row(iid)
            self._auto_save()
            return

        if col_key == "startnummer":
            try:
                new_sn = int(new_text)
            except ValueError:
                messagebox.showwarning("Ungültige Startnummer", "Die Startnummer muss eine ganze Zahl sein.")
                return
            for oid in self.data:
                if oid != iid and self.data[oid].get("startnummer") == new_sn:
                    messagebox.showwarning("Doppelte Startnummer", f"Startnummer {new_sn} ist bereits vergeben.")
                    return
            self.data[iid]["startnummer"] = new_sn
            self._update_tree_row(iid)
            self._auto_save()
            return

        # Disziplin-Ergebnis?
        if col_key.endswith("_erg"):
            try:
                new_val = self._parse_float(new_text, allow_empty=False)
            except ValueError:
                messagebox.showwarning("Ungültige Eingabe", "Bitte eine Zahl eingeben.")
                return
            self.data[iid][col_key] = float(new_val)
            # Zeile neu berechnen + Ränge aktualisieren
            self._recalc_row(iid)
            self._update_tree_row(iid)
            self._recalc_ranks_and_update()
            self._auto_save()
            return

    def _cancel_inline_edit(self, event=None):
        if self._edit_entry:
            self._edit_entry.destroy()
        self._edit_entry = None
        self._edit_iid_col = None

    # =========================
    # Sortierung
    # =========================
    def _is_numeric_column(self, key):
        return key in self.numeric_columns

    def on_sort_column(self, col):
        asc = self.sort_state.get(col, True)
        asc = not asc
        self.sort_state[col] = asc
        self._apply_sort(col, asc)
        self._last_sort_col = col

    def _apply_sort(self, col, ascending=True):
        children = list(self.tree.get_children(""))

        def get_val(iid):
            row = self.data.get(iid, {})
            v = row.get(col)
            if v is None and col in ("name",):
                return str(row.get("name") or "")
            if v is None:
                try:
                    idx = self.tree["displaycolumns"].index(col)
                    val = self.tree.item(iid, "values")[idx]
                except Exception:
                    val = ""
                if self._is_numeric_column(col):
                    try:
                        return float(str(val).replace(",", "."))
                    except ValueError:
                        return float("-inf")
                return str(val)
            if self._is_numeric_column(col):
                try:
                    return float(v)
                except (TypeError, ValueError):
                    return float("-inf")
            return str(v)

        sorted_children = sorted(children, key=get_val, reverse=not ascending)
        for idx, iid in enumerate(sorted_children):
            self.tree.move(iid, "", idx)

    # =========================
    # Export CSV
    # =========================
    def export_csv(self):
        import csv
        import datetime
        filename = filedialog.asksaveasfilename(
            title="CSV speichern",
            defaultextension=".csv",
            filetypes=[("CSV-Datei", "*.csv")]
        )
        if not filename:
            return

        # Alle Spalten exportieren, einschließlich inaktiver Disziplinen
        cols = self.all_columns
        headers = [self.column_headers[c] for c in cols]

        with open(filename, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f, delimiter=";")
            # Titel und Datum hinzufügen
            w.writerow(["Titel", self.ent_title.get()])
            w.writerow(["Datum", datetime.date.today().isoformat()])
            w.writerow([])  # Leere Zeile
            w.writerow(headers)
            for iid in self.tree.get_children():
                row = self.data[iid]
                out = []
                for c in cols:
                    out.append(row[c] if c in ("name",) else self._format_number(row.get(c)))
                w.writerow(out)

        messagebox.showinfo("Export erfolgreich", f"Die CSV wurde gespeichert:\n{filename}")

    # =========================
    # Auto-Save CSV
    # =========================
    def _auto_save(self):
        import csv
        import datetime
        import os
        # Dateiname aus Titel erstellen, ungültige Zeichen ersetzen
        title = self.ent_title.get().strip() or "Wettbewerb"
        safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = os.path.join(os.getcwd(), safe_title + ".csv")

        # Alle Spalten exportieren, einschließlich inaktiver Disziplinen
        cols = self.all_columns
        headers = [self.column_headers[c] for c in cols]

        try:
            with open(filename, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f, delimiter=";")
                # Titel und Datum hinzufügen
                w.writerow(["Titel", self.ent_title.get()])
                w.writerow(["Datum", datetime.date.today().isoformat()])
                w.writerow([])  # Leere Zeile
                w.writerow(headers)
                for iid in self.tree.get_children():
                    row = self.data[iid]
                    out = []
                    for c in cols:
                        out.append(row[c] if c in ("name",) else self._format_number(row.get(c)))
                    w.writerow(out)
        except Exception as e:
            # Optional: Fehlermeldung, aber da automatisch, still ignorieren
            pass

    # =========================
    # Load CSV
    # =========================
    def load_csv(self):
        import csv
        filename = filedialog.askopenfilename(
            title="CSV laden",
            filetypes=[("CSV-Datei", "*.csv"), ("Alle Dateien", "*")]
        )
        if not filename:
            return

        try:
            with open(filename, newline="", encoding="utf-8") as f:
                reader = csv.reader(f, delimiter=";")
                rows = list(reader)
        except Exception as e:
            messagebox.showerror("Fehler beim Einlesen", f"CSV konnte nicht eingelesen werden:\n{e}")
            return

        if len(rows) < 4:
            messagebox.showerror("Fehler", "CSV-Datei ist zu kurz oder hat nicht das erwartete Format.")
            return

        # Erste Zeile: Titel
        if len(rows[0]) >= 2 and rows[0][0] == "Titel":
            self.ent_title.delete(0, tk.END)
            self.ent_title.insert(0, rows[0][1])

        # Zweite Zeile: Datum (ignorieren)
        # Dritte Zeile: leer (ignorieren)

        # Vierte Zeile: Headers
        headers = rows[3]
        if not headers:
            messagebox.showerror("Fehler", "CSV-Datei benötigt Kopfzeile.")
            return

        # Datenzeilen ab Zeile 4
        csv_rows = []
        for row in rows[4:]:
            if row:  # Überspringe leere Zeilen
                csv_rows.append(dict(zip(headers, row)))

        def _get(r, *keys):
            for k in keys:
                v = r.get(k)
                if v is not None and str(v).strip() != "":
                    return str(v).strip()
            return None

        def _normalize_key(k):
            return k.strip().lower() if k is not None else ""

        self.data.clear()

        for idx, raw_row in enumerate(csv_rows, start=1):
            row = {k: (v.strip() if isinstance(v, str) else v) for k, v in raw_row.items()}
            normalized = {(_normalize_key(k)): v for k, v in row.items()}

            name = _get(normalized, "name", "participant", "teilnehmer")
            if not name:
                continue

            startnr = None
            for c in ["startnummer", "startnr", "startNr", "start", "Startnr."]:
                if c.lower() in normalized and normalized[c.lower()] not in (None, ""):
                    try:
                        startnr = int(float(normalized[c.lower()]))
                        break
                    except (ValueError, TypeError):
                        startnr = None

            item = {
                "name": name,
                "startnummer": startnr,
                "gesamt": None,
                "gesamtrang": None,
            }

            for d in DISCIPLINES:
                keys = [
                    f"{d.label} erg".lower(),
                    f"{d.code}_erg".lower(),
                    f"{d.label.lower()}_erg", 
                    f"{d.label}erg".lower(),
                ]
                val = None
                for k in keys:
                    if k in normalized and normalized[k] not in (None, ""):
                        try:
                            val = self._parse_float(normalized[k], allow_empty=False)
                        except Exception:
                            val = None
                        break
                item[f"{d.code}_erg"] = float(val) if val is not None else 0.0
                item[f"{d.code}_pkt"] = None
                item[f"{d.code}_rang"] = None

            self.data[str(idx)] = item

        # Rebuild and recalc results
        self._rebuild_dynamic_ui_and_tree()
        for iid in self.data.keys():
            self._recalc_row(iid)
        self._recalc_ranks_and_update()

        messagebox.showinfo("Import erfolgreich", f"{len(self.data)} Zeilen aus '{os.path.basename(filename)}' geladen.")

    # =========================
    # Export PDF Gesamtliste
    # =========================
    def export_pdf(self):
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
        except Exception:
            messagebox.showerror(
                "ReportLab fehlt",
                "Für den PDF-Export wird das Paket 'reportlab' benötigt.\n"
                "Installiere es mit:\n\n    pip install reportlab"
            )
            return

        filename = filedialog.asksaveasfilename(
            title="PDF speichern",
            defaultextension=".pdf",
            filetypes=[("PDF-Datei", "*.pdf")]
        )
        if not filename:
            return

        doc = SimpleDocTemplate(
            filename,
            pagesize=landscape(A4),
            leftMargin=15*mm,
            rightMargin=15*mm,
            topMargin=15*mm,
            bottomMargin=15*mm,
        )

        styles = getSampleStyleSheet()
        story = []

        title_text = self.ent_title.get().strip() or "Wettbewerb"
        story.append(Paragraph(title_text, styles["Title"]))
        story.append(Spacer(1, 6))

        # Tabelle: Basis + aktive Disziplinspalten (Erg/Pkt/Rang)
        headers = ["Name", "Startnr.", "Gesamt", "Gesamtrang"]
        col_keys = ["name", "startnummer", "gesamt", "gesamtrang"]
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                headers += [f"{d.label} Erg", f"{d.label} Pkt", f"{d.label} Rang"]
                col_keys += [f"{d.code}_erg", f"{d.code}_pkt", f"{d.code}_rang"]

        data_rows = []
        for iid in self.tree.get_children():
            r = self.data[iid]
            row_vals = []
            for k in col_keys:
                if k == "name":
                    row_vals.append(str(r["name"]))
                else:
                    row_vals.append(self._format_number(r.get(k)))
            data_rows.append(row_vals)

        table_data = [headers] + data_rows

        # Spaltenbreiten heuristisch
        # Basis: 60 + 18 + 22 + 24 = 124mm, Rest auf Disziplinen; pro Disziplin ~ (22+22+20)=64mm
        col_widths = []
        base_widths = [45*mm, 11*mm, 11*mm, 11*mm]
        col_widths.extend(base_widths)
        for d in DISCIPLINES:
            if self.disc_state[d.code].get():
                col_widths.extend([11*mm, 11*mm, 9*mm])

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 5),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),

            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 6),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("ALIGN", (0, 1), (0, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),

            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ]))
        story.append(tbl)

        try:
            doc.build(story)
        except Exception as e:
            messagebox.showerror("PDF-Fehler", f"Beim Erstellen des PDFs ist ein Fehler aufgetreten:\n{e}")
            return

        messagebox.showinfo("Export erfolgreich", f"Die PDF wurde gespeichert:\n{filename}")

    # =========================
    # Export: Individuelle Berichte (A4, kompakte Disziplin-Tabelle, Logo)
    # =========================
    def export_individual_reports(self):
        # --- Bibliotheken prüfen ---
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        except Exception:
            pdf_available = False
        else:
            pdf_available = True
    
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except Exception:
            docx_available = False
        else:
            docx_available = True
    
        if not pdf_available and not docx_available:
            messagebox.showerror(
                "Fehlende Pakete",
                "Weder ReportLab (PDF) noch python-docx (Word) sind installiert.\n"
                "Installiere mindestens eines davon:\n\n"
                "pip install reportlab\n"
                "pip install python-docx"
            )
            return
    
        if not self.data:
            messagebox.showwarning("Keine Daten", "Es sind keine Teilnehmer vorhanden.")
            return
    
        # Ränge aktualisieren
        self._recalc_ranks_and_update()
    
        # Sortierung
        entries = sorted(
            self.data.items(),
            key=lambda kv: ((kv[1].get("gesamtrang") or 10**9), str(kv[1].get("name") or ""))
        )
    
        # --- Dateidialog: Nutzer entscheidet Format ---
        out_file = filedialog.asksaveasfilename(
            title="Bericht speichern",
            defaultextension=".pdf",
            filetypes=[
                ("PDF-Datei", "*.pdf"),
                ("Word-Dokument", "*.docx")
            ]
        )
        if not out_file:
            return
    
        # --- Format bestimmen ---
        is_pdf = out_file.lower().endswith(".pdf")
        is_docx = out_file.lower().endswith(".docx")
    
        title_text = self.ent_title.get().strip() or "Wettbewerb"
    
        # ---------------------------------------------------------
        # 1) WORD EXPORT
        # ---------------------------------------------------------
        if is_docx:
            if not docx_available:
                messagebox.showerror("Fehler", "python-docx ist nicht installiert.")
                return
    
            doc = Document()
    
            def add_logo(document):
                if not self.logo_path:
                    return
                try:
                    document.add_picture(self.logo_path, width=Inches(2.0))
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
    
            for idx, (iid, row) in enumerate(entries):
                h = doc.add_heading(title_text, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
                h2 = doc.add_heading("Overall award", level=2)
                h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
                add_logo(doc)
                doc.add_paragraph("")
    
                # Teilnehmerkopf
                table = doc.add_table(rows=0, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
                def add_info(label, value):
                    r = table.add_row().cells
                    r[0].text = label
                    r[1].text = str(value)
    
                add_info("Name:", row.get("name") or "")
                add_info("Gesamtpunkte:", self._format_number(row.get("gesamt")))
                add_info("Gesamtrang:", self._format_number(row.get("gesamtrang")))
    
                doc.add_paragraph("")
    
                # Disziplin-Tabelle
                disc_table = doc.add_table(rows=1, cols=4)
                disc_table.style = "Table Grid"
                hdr = disc_table.rows[0].cells
                hdr[0].text = "Disziplin"
                hdr[1].text = "Ergebnis"
                hdr[2].text = "Punkte"
                hdr[3].text = "Rang"
    
                for d in DISCIPLINES:
                    if not self.disc_state[d.code].get():
                        continue
                    row_cells = disc_table.add_row().cells
                    row_cells[0].text = d.label
                    row_cells[1].text = self._format_number(row.get(f"{d.code}_erg"))
                    row_cells[2].text = self._format_number(row.get(f"{d.code}_pkt"))
                    row_cells[3].text = self._format_number(row.get(f"{d.code}_rang"))
    
                if idx < len(entries) - 1:
                    doc.add_page_break()
    
            try:
                doc.save(out_file)
            except Exception as e:
                messagebox.showerror("Word-Fehler", f"Fehler beim Speichern:\n{e}")
                return
    
            messagebox.showinfo("Export erfolgreich", f"Word-Dokument gespeichert:\n{out_file}")
            return
    
        # ---------------------------------------------------------
        # 2) PDF EXPORT (dein bestehender Code)
        # ---------------------------------------------------------
        if is_pdf:
            if not pdf_available:
                messagebox.showerror("Fehler", "ReportLab ist nicht installiert.")
                return
    
    
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25, leading=24, spaceAfter=6)
            h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=20, spaceBefore=6, spaceAfter=6, alignment=1)
            label_style = ParagraphStyle("Label", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=11)
            text_style = ParagraphStyle("Text", parent=styles["Normal"], fontSize=11)
            name_style = ParagraphStyle("Name", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=15, spaceBefore=6, spaceAfter=6, alignment=1)
            rank_style = ParagraphStyle("Rank", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=20, spaceBefore=6, spaceAfter=6, alignment=1)
            points_style = ParagraphStyle("Points", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, spaceBefore=6, spaceAfter=6, alignment=1)
    
            def make_logo():
                if not self.logo_path:
                    return None
                try:
                    img = Image(self.logo_path)
                    max_w, max_h = 160*mm, 160*mm
                    iw, ih = img.imageWidth, img.imageHeight
                    scale = min(max_w/iw, max_h/ih)
                    img.drawWidth = iw * scale
                    img.drawHeight = ih * scale
                    return img
                except Exception:
                    return None
    
            title_text = self.ent_title.get().strip() or "Wettbewerb"
    
            def build_story_for_row(row):
                story = []
                logo = make_logo()
                story.append(Paragraph(title_text, title_style))
                story.append(Spacer(1, 6))
                story.append(Paragraph('<para alignment="center">Overall award</para>', h2_style))
                story.append(Spacer(1, 6))
                
                if logo:
                    logo.hAlign = "CENTER"
                    story.append(logo)
                
                story.append(Spacer(1, 20))

                _name = str(row.get("name") or "")
                _gesamt = self._format_number(row.get("gesamt"))
                _rank = self._format_number(row.get("gesamtrang"))
                
                story.append(Paragraph(f"{_rank}.Place", rank_style))
                story.append(Spacer(1, 4))
                story.append(Paragraph(f"with {_gesamt} points", points_style))
                story.append(Spacer(1, 4))
                story.append(Paragraph(f"{_name}", name_style))
#                # Teilnehmerkopf
#                info_tbl = Table([
#                    [Paragraph("Name:", label_style), Paragraph(str(row.get("name") or ""), text_style)],
#                    #[Paragraph("Startnummer:", label_style), Paragraph(self._format_number(row.get("startnummer")), text_style)],
#                    [Paragraph("Gesamtpunkte:", label_style), Paragraph(self._format_number(row.get("gesamt")), text_style)],
#                    [Paragraph("Gesamtrang:", label_style), Paragraph(self._format_number(row.get("gesamtrang")), text_style)]
#                ], colWidths=[40*mm, None])
#            #    info_tbl.setStyle(TableStyle([
#            #        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
#            #        ("LEFTPADDING", (0, 0), (-1, -1), 0),
#            #        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
#            #    ]))
#                
#                info_tbl.setStyle(TableStyle([
#                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
#                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
#                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
#                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
#                    ("TOPPADDING", (0, 0), (-1, -1), 4),
#                ]))
#                info_tbl.hAlign = "CENTER"
#          
#                story.append(info_tbl)
                story.append(Spacer(1, 30))
    
                # Kompakte Disziplin-Tabelle (nur aktive)
                tbl_headers = ["Disziplin", "Ergebnis", "Punkte", "Rang"]
                tbl_rows = []
                for d in DISCIPLINES:
                    if not self.disc_state[d.code].get():
                        continue
                    tbl_rows.append([
                        d.label,
                        self._format_number(row.get(f"{d.code}_erg")),
                        self._format_number(row.get(f"{d.code}_pkt")),
                        self._format_number(row.get(f"{d.code}_rang")),
                    ])
                table_data = [tbl_headers] + tbl_rows
                disc_tbl = Table(table_data, colWidths=[28*mm, 28*mm, 28*mm, 22*mm])
                disc_tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
    
                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 1), (-1, -1), 10),
                    ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                    ("ALIGN", (0, 1), (0, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ]))
                story.append(disc_tbl)
                return story
    
            from reportlab.platypus import PageBreak
            doc = SimpleDocTemplate(out_file, pagesize=A4, leftMargin=18*mm, rightMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm)
            story = []
            for idx, (iid, row) in enumerate(entries):
                story.extend(build_story_for_row(row))
                if idx < len(entries) - 1:
                    story.append(PageBreak())
            try:
                doc.build(story)
            except Exception as e:
                messagebox.showerror("PDF-Fehler", f"Beim Erstellen des PDFs ist ein Fehler aufgetreten:\n{e}")
                return
    
            messagebox.showinfo("Export erfolgreich", f"PDF gespeichert:\n{out_file}")
            return
    
    
    # =========================
    # Hilfsfunktion: Spalten basierend auf Event definieren
    # =========================
    def _get_scoresheet_columns(self, event):
        """
        Gibt die Spalten für das Scoresheet basierend auf dem Event zurück.
        Rückgabe:
          - header_rows: Liste von Header-Zeilen (für Gruppierung/Unterspalten)
          - num_empty_cols: Anzahl der leeren Datenspalten zwischen Name und Result (für die Dateneinträge)
          - col_widths: Liste der Spaltenbreiten in mm
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm

        if event == "ACC":
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 11)] + ["Result"]
            num_empty_cols = 10
            col_widths = [12*mm, 40*mm] + [12*mm] * 10 + [15*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "AUS":
            main_headers = ["Startnr", "Name"]
            sub_headers = ["", ""]
            for i in range(1, 6):
                main_headers.extend([f"Throw {i}", "", "", ""])
                sub_headers.extend(["Dist", "Catch", "Acc", "To"])
            main_headers.append("Result")
            sub_headers.append("")
            num_empty_cols = 20
            col_widths = [12*mm, 35*mm] + [10*mm] * 20 + [15*mm]
            return [main_headers, sub_headers], num_empty_cols, col_widths
        elif event == "END":
            distances = [5, 10, 15, 20, 25, 30, 35, 40, 45, 50, 55, 60, 65, 70, 75, 80]
            headers = ["Startnr", "Name"] + [str(d) for d in distances] + ["Result"]
            num_empty_cols = 16
            col_widths = [12*mm, 35*mm] + [10*mm] * 16 + [15*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "FC":
            headers = ["Startnr", "Name", "Round 1", "Round 2", "Result"]
            num_empty_cols = 2
            col_widths = [12*mm, 40*mm] + [40*mm] * 2 + [40*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "MTA":
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 6)] + ["Result"]
            num_empty_cols = 5
            col_widths = [12*mm, 40*mm] + [20*mm] * 5 + [20*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "TC":
            main_headers = ["Startnr", "Name"] + ["Throw 1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "", "Doubl.1", "", "2", "", "3", "", "4", "", "5", ""] + ["Result"]
            sub_headers = ["", ""] + ["Left-\nhand\nclean", "Right-\nhand\nclean", "2 hand\nbehind\nback", "2 hand\nunder\nthe-leg", "Eagle\ncatch", "Hacky\ncatch", "Tunnel\ncatch", "1 hand\nbehind\nback", "1 hand\nunder\nthe leg", "Foot\ncatch", "Total\nSingle", "2 hand\nbehind-\nthe-back", "2 hand\nunder-\nthe-leg", "Left-\nhand\nclean", "Hacky\ncatch", "Right-\nhand\nclean", "Tunnel\ncatch", "1 hand\nbehind\nthe back", "1 hand\nunder\nthe leg", "Eagle\ncatch", "Foot\ncatch"] + [""]
            points_headers = ["", ""] + ["(3)", "(3)", "(4)", "(3)", "(4)", "(7)", "(5)", "(7)", "(6)", "(8)", "", "(4)", "(3)", "(3)", "(7)", "(3)", "(5)", "(7)", "(6)", "(4)", "(8)"] + [""]
            num_empty_cols = 21
            col_widths = [12*mm, 35*mm] + [9*mm] * 21 + [15*mm]
            return [main_headers, sub_headers, points_headers], num_empty_cols, col_widths
        elif event == "TAPIR":
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 6)] + ["Result"]
            num_empty_cols = 5
            col_widths = [12*mm, 40*mm] + [15*mm] * 5 + [15*mm]
            return [headers], num_empty_cols, col_widths
        elif event == "TIMED":
            headers = ["Startnr", "Name", "Round 1", "Round 2", "Result"]
            num_empty_cols = 2
            col_widths = [12*mm, 40*mm] + [40*mm] * 2 + [40*mm]
            return [headers], num_empty_cols, col_widths
        else:
            headers = ["Startnr", "Name"] + [f"{i}. Throw" for i in range(1, 11)] + ["Result"]
            num_empty_cols = 10
            col_widths = [12*mm, 30*mm] + [12*mm] * 10 + [15*mm]
            return [headers], num_empty_cols, col_widths
    
    # =========================
    # Export Scoresheet (DOCX/PDF mit Circles)
    # =========================
    def export_scoresheet(self):
        # --- Bibliotheken prüfen ---
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet
        except Exception:
            pdf_available = False
        else:
            pdf_available = True
    
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.section import WD_ORIENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except Exception:
            docx_available = False
        else:
            docx_available = True
    
        if not pdf_available and not docx_available:
            messagebox.showerror(
                "Fehlende Pakete",
                "Weder ReportLab (PDF) noch python-docx (DOCX) sind installiert.\n"
                "Installiere mindestens eines davon:\n\n"
                "pip install reportlab\n"
                "pip install python-docx"
            )
            return
    
        if not self.data:
            messagebox.showwarning("Keine Daten", "Es sind keine Teilnehmer vorhanden.")
            return
    
        # --- Eingaben auslesen ---
        try:
            num_circles = int(self.ent_circle.get().strip() or "1")
            if num_circles < 1:
                num_circles = 1
        except (ValueError, AttributeError):
            num_circles = 1
    
        sort_method = self.sheetsort_var.get() if hasattr(self, 'sheetsort_var') else "StartNr"
        event = self.event_var.get() if hasattr(self, 'event_var') else "ACC"
        
        # Ränge aktualisieren
        self._recalc_ranks_and_update()
    
        # --- Daten sammeln und sortieren ---
        if sort_method == "Rank":
            # Sortierung nach Rang
            entries = sorted(
                self.data.items(),
                key=lambda kv: ((kv[1].get("gesamtrang") or 10**9), str(kv[1].get("name") or ""))
            )
        else:
            # Sortierung nach Startnummer (Standard)
            entries = sorted(
                self.data.items(),
                key=lambda kv: (int(kv[1].get("startnummer") or 999), str(kv[1].get("name") or ""))
            )
    
        # --- Daten in Circles aufteilen ---
        circles_data = self._distribute_entries_to_circles(entries, num_circles, sort_method)
    
        # --- Spalten für diesen Event ---
        header_rows, num_empty_cols, col_widths = self._get_scoresheet_columns(event)
        num_cols = len(header_rows[-1])

        # --- Dialog: Format wählen ---
        out_file = filedialog.asksaveasfilename(
            title="Scoresheet speichern",
            defaultextension=".docx",
            filetypes=[
                ("Word-Dokument", "*.docx"),
                ("PDF-Datei", "*.pdf")
            ]
        )
        if not out_file:
            return
    
        is_docx = out_file.lower().endswith(".docx")
        is_pdf = out_file.lower().endswith(".pdf")
    
        title_text = self.ent_title.get().strip() or "Wettbewerb"
    
        # ---------------------------------------------------------
        # 1) DOCX EXPORT
        # ---------------------------------------------------------
        if is_docx:
            if not docx_available:
                messagebox.showerror("Fehler", "python-docx ist nicht installiert.")
                return
    
            doc = Document()
            
            # Setze Querformat (Landscape)
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            
            def set_cell_background(cell, fill):
                """Hintergrundfarbe einer Zelle setzen"""
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), fill)
                cell._element.get_or_add_tcPr().append(shading_elm)
    
            for circle_idx, circle_entries in enumerate(circles_data, 1):
                if circle_idx > 1:
                    doc.add_page_break()
    
                # Titel für Circle
                heading = doc.add_heading(f"{title_text} - {event} - Circle {circle_idx}", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
# Tabelle erstellen (Header-Zeilen berücksichtigen)
                table = doc.add_table(rows=len(header_rows), cols=num_cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header (mehrzeilig möglich)
                for header_row_idx, header_row in enumerate(header_rows):
                    hdr_cells = table.rows[header_row_idx].cells
                    for col_idx, header_text in enumerate(header_row):
                        hdr_cells[col_idx].text = header_text
                        set_cell_background(hdr_cells[col_idx], "D3D3D3")
                        for paragraph in hdr_cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(9)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
                # Datenzeilen
                for iid, row in circle_entries:
                    row_cells = table.add_row().cells
                    startnr = row.get("startnummer", "")
                    name = row.get("name", "")

                    row_cells[0].text = str(startnr)
                    row_cells[1].text = str(name)
                    # Leere Zellen für Einträge (Result bleibt leer)
                    for i in range(2, num_cols):
                        row_cells[i].text = ""
    
                    # Zellausrichtung
                    for idx, cell in enumerate(row_cells):
                        for paragraph in cell.paragraphs:
                            if idx == 0 or idx == num_cols - 1:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif idx == 1:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
                doc.add_paragraph("")
    
            try:
                doc.save(out_file)
            except Exception as e:
                messagebox.showerror("DOCX-Fehler", f"Fehler beim Speichern:\n{e}")
                return
    
            messagebox.showinfo("Export erfolgreich", f"Scoresheet (DOCX) gespeichert:\n{out_file}")
            return
    
        # ---------------------------------------------------------
        # 2) PDF EXPORT
        # ---------------------------------------------------------
        if is_pdf:
            if not pdf_available:
                messagebox.showerror("Fehler", "ReportLab ist nicht installiert.")
                return
    
            styles = getSampleStyleSheet()
            doc = SimpleDocTemplate(
                out_file,
                pagesize=landscape(A4),
                leftMargin=10*mm,
                rightMargin=10*mm,
                topMargin=10*mm,
                bottomMargin=10*mm,
            )
            story = []
    
            for circle_idx, circle_entries in enumerate(circles_data, 1):
                if circle_idx > 1:
                    story.append(PageBreak())

                # Titel für Circle
                story.append(Paragraph(f"{title_text} - {event} - Circle {circle_idx}", styles["Title"]))
                story.append(Spacer(1, 6))

                table_data = [list(r) for r in header_rows]

                for iid, row in circle_entries:
                    startnr = str(row.get("startnummer", ""))
                    name = str(row.get("name", ""))

                    row_vals = [startnr, name]
                    row_vals.extend([""] * (num_empty_cols + 1))  # Leere Zellen für Einträge und Result
                    table_data.append(row_vals)

                tbl = Table(table_data, colWidths=col_widths, repeatRows=len(header_rows))
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d3d3d3")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 7),
                    ("TOPPADDING", (0, 0), (-1, 0), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 5),
    
                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 1), (-1, -1), 7),
                    ("ALIGN", (0, 1), (0, -1), "CENTER"),
                    ("ALIGN", (1, 1), (1, -1), "LEFT"),
                    ("ALIGN", (2, 1), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 12))
    
            try:
                doc.build(story)
            except Exception as e:
                messagebox.showerror("PDF-Fehler", f"Beim Erstellen des PDFs ist ein Fehler aufgetreten:\n{e}")
                return
    
            messagebox.showinfo("Export erfolgreich", f"Scoresheet (PDF) gespeichert:\n{out_file}")
    
    # =========================
    # Hilfsfunktion: Entries auf Circles verteilen
    # =========================
    def _distribute_entries_to_circles(self, entries, num_circles, sort_method):
        """
        Verteilt Einträge auf mehrere Circles.
        - sort_method == "StartNr": sequentielle Verteilung nach Startnummer
        - sort_method == "Rank": zyklische Verteilung nach Rang
        """
        circles_data = [[] for _ in range(num_circles)]
        
        if sort_method == "Rank":
            # Zyklische Verteilung nach Rang (Rund-Robin)
            # Rang 1 → Circle 0, Rang 2 → Circle 1, ..., Rang n → Circle (n-1) % num_circles
            for idx, entry in enumerate(entries):
                circle_idx = idx % num_circles
                circles_data[circle_idx].append(entry)
        else:
            # Sequentielle Verteilung nach Startnummer
            total_entries = len(entries)
            entries_per_circle = max(1, total_entries // num_circles)
            
            for idx, entry in enumerate(entries):
                # Berechne Circle-Index basierend auf Position
                circle_idx = min(idx // entries_per_circle, num_circles - 1)
                circles_data[circle_idx].append(entry)
        
        return circles_data

    # =========================
    # Start-/Hilfsfunktionen
    # =========================
    def _is_active(self, code):
        return bool(self.disc_state[code].get())


# ==============================
# Start
# ==============================
if __name__ == "__main__":
    app = ScoreTableApp()
    app.mainloop()

'''spaltenbreite für scoresheets bei header def mit übergeben - erl.
        - spaltenbreiten definieren, mehr Höhe für die Zeilen
   disziplin tapir hinzufügen ->erledigt
   scoresheet für tapir anpassen
   formel für Tapir anpassen
   Tickfang scoresheet mit Wurf 1... erweitern wie bei AUS
    
    
   
 
'''
